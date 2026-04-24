from __future__ import annotations

import base64
from datetime import datetime
import difflib
from io import BytesIO
import json
import os
from pathlib import Path
import re
import shutil
from typing import Any
from urllib import error as urllib_error
from urllib import parse as urllib_parse
from urllib import request as urllib_request

import pandas as pd
import streamlit as st
try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError:  # pragma: no cover - runtime fallback
    Document = None
    WD_SECTION = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    Inches = None
    Pt = None
try:
    import openpyxl  # noqa: F401
except ImportError:  # pragma: no cover - runtime fallback
    openpyxl = None

from config import settings
from graph import (
    ARCHITECT_APPROVED,
    ARCHITECT_REJECTED,
    PENDING_ARCHITECT_APPROVAL,
    PENDING_SME_APPROVAL,
    SME_APPROVED,
    SME_REJECTED,
    run_workflow,
    stream_resume_from_requirements_approval,
    stream_resume_from_technical_spec_approval,
    stream_workflow,
)
from utils.cache import AgentCache
from agents import ai_code_review as ai_code_review_agent
from agents import data_mapping as data_mapping_agent
from agents import validation as validation_agent


st.set_page_config(
    page_title="AI Enabled App Migration Platform",
    page_icon="AI",
    layout="wide",
    initial_sidebar_state="expanded",
)

MODEL_OPTIONS = [
    "gpt-5.4-mini",
    "gpt-5.4",
    "gpt-5-mini",
    "gpt-5",
    "gpt-4o-mini",
    "claude-sonnet-4-6",
    "claude-opus-4-1-20250805",
]

MODEL_LABELS = {
    "gpt-5.4-mini": "GPT-5.4 Mini",
    "gpt-5.4": "GPT-5.4",
    "gpt-5-mini": "GPT-5 Mini",
    "gpt-5": "GPT-5",
    "gpt-4o-mini": "GPT-4o Mini",
    "claude-sonnet-4-6": "Claude Sonnet 4.6",
    "claude-opus-4-1-20250805": "Claude Opus 4.1",
}

WORKFLOW_PHASES = [
    ("reverse_legacy", "Analyzing legacy VB and SQL files"),
    ("collate_legacy", "Collating legacy business flow"),
    ("reverse_target", "Analyzing target code and SQL files"),
    ("collate_target", "Collating target business flow"),
    ("gap", "Analyzing requirement gaps between legacy and target system"),
    ("requirements", "Drafting incremental business requirements"),
    ("technical_spec", "Drafting incremental technical specification"),
    ("forward_engineering", "Generating forward-engineered code components"),
]

DASHBOARD_STAGES = [
    {"key": "reverse_legacy", "label": "Legacy Reverse", "result_keys": ["legacy_code_reverse_spec", "legacy_sql_reverse_spec"], "detail_section": "Step Outputs"},
    {"key": "collate_legacy", "label": "Legacy Collate", "result_keys": ["legacy_spec"], "detail_section": "Legacy Spec"},
    {"key": "reverse_target", "label": "Target Reverse", "result_keys": ["target_code_reverse_spec", "target_sql_reverse_spec"], "detail_section": "Step Outputs"},
    {"key": "collate_target", "label": "Target Collate", "result_keys": ["target_spec", "collated_spec"], "detail_section": "Target Spec"},
    {"key": "gap", "label": "Gap Analysis", "result_keys": ["gap_analysis"], "detail_section": "Gap Analysis"},
    {"key": "requirements", "label": "Requirements Draft", "result_keys": ["requirements_draft"], "detail_section": "Requirements Draft"},
    {"key": "requirements_approval", "label": "SME Approval", "approval": "requirements"},
    {"key": "technical_spec", "label": "Technical Spec", "result_keys": ["technical_spec_draft"], "detail_section": "Technical Specification Draft"},
    {"key": "technical_spec_approval", "label": "Architect Approval", "approval": "technical_spec"},
    {"key": "forward_engineering", "label": "Forward Engineering", "result_keys": ["forward_engineering_output"], "detail_section": "Forward Engineering"},
    {"key": "forward_engineering_proof", "label": "FE Proof", "result_keys": ["forward_engineering_output"], "detail_section": "Forward Engineering Proof"},
    {"key": "data_mapping", "label": "Final Data Mapping", "result_keys": ["data_mapping_result"], "detail_section": "Final Data Mapping"},
    {"key": "validation", "label": "Validation", "result_keys": ["validation_result"], "detail_section": "Validation"},
    {"key": "ai_code_review", "label": "AI Code Review", "result_keys": ["ai_code_review_result"], "detail_section": "AI Code Review"},
]


def apply_enterprise_theme() -> None:
    st.markdown(
        """
        <style>
        .stApp {
            background:
                radial-gradient(circle at top right, rgba(11, 78, 140, 0.14), transparent 28%),
                linear-gradient(180deg, #f4f7fb 0%, #edf2f7 100%);
        }
        .block-container {
            padding-top: 1.25rem;
            padding-bottom: 2rem;
            max-width: 1400px;
        }
        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, #0f172a 0%, #16263d 100%);
            border-right: 1px solid rgba(255, 255, 255, 0.08);
        }
        [data-testid="stSidebar"] * {
            color: #e5eef8;
        }
        [data-testid="stSidebar"] p,
        [data-testid="stSidebar"] span,
        [data-testid="stSidebar"] div,
        [data-testid="stSidebar"] label {
            color: #e5eef8;
        }
        [data-testid="stSidebar"] input,
        [data-testid="stSidebar"] textarea,
        [data-testid="stSidebar"] [data-baseweb="input"] input {
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            background-color: #ffffff !important;
        }
        [data-testid="stSidebar"] [data-baseweb="select"] > div,
        [data-testid="stSidebar"] [data-baseweb="select"] span,
        [data-testid="stSidebar"] [data-baseweb="popover"] *,
        [data-testid="stSidebar"] [data-baseweb="base-input"],
        [data-testid="stSidebar"] [role="combobox"],
        [data-testid="stSidebar"] [role="listbox"],
        [data-testid="stSidebar"] [role="option"] {
            background-color: #ffffff !important;
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
        }
        [data-testid="stSidebar"] label,
        [data-testid="stSidebar"] .stMarkdown,
        [data-testid="stSidebar"] .stCaption,
        [data-testid="stSidebar"] .stText,
        [data-testid="stSidebar"] .stSelectbox label,
        [data-testid="stSidebar"] .stToggle label {
            color: #e5eef8 !important;
        }
        .hero-card, .panel-card, .metric-card, .gap-card {
            background: rgba(255, 255, 255, 0.92);
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: 18px;
            box-shadow: 0 18px 38px rgba(15, 23, 42, 0.08);
        }
        .hero-card {
            padding: 1.5rem 1.75rem;
            min-height: 180px;
        }
        .panel-card {
            padding: 1.1rem 1.2rem;
        }
        .phase-card {
            height: 206px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
        }
        .phase-title {
            min-height: 82px;
            display: flex;
            align-items: flex-start;
        }
        .metric-card {
            padding: 1rem 1.1rem;
        }
        .gap-card {
            padding: 1rem 1.15rem;
            height: 100%;
        }
        .eyebrow {
            text-transform: uppercase;
            letter-spacing: 0.08em;
            font-size: 0.72rem;
            color: #0b4e8c;
            font-weight: 700;
        }
        .hero-title {
            font-size: 2.15rem;
            line-height: 1.05;
            margin: 0.35rem 0 0.55rem 0;
            color: #10233b;
            font-weight: 800;
        }
        .hero-copy, .muted-copy {
            color: #4b5d73;
            font-size: 0.98rem;
        }
        .section-title {
            font-size: 1.1rem;
            font-weight: 700;
            color: #10233b;
            margin-bottom: 0.65rem;
        }
        .status-chip {
            display: inline-block;
            padding: 0.35rem 0.65rem;
            border-radius: 999px;
            background: transparent;
            color: #0b4e8c;
            font-size: 0.8rem;
            font-weight: 700;
            margin-right: 0.4rem;
            margin-bottom: 0.4rem;
        }
        .risk-high {
            border-left: 5px solid #c62828;
        }
        .risk-medium {
            border-left: 5px solid #ef6c00;
        }
        .risk-low {
            border-left: 5px solid #2e7d32;
        }
        .small-label {
            font-size: 0.78rem;
            color: #6b7c93;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            margin-bottom: 0.3rem;
        }
        .metric-value {
            font-size: 1.8rem;
            font-weight: 800;
            color: #10233b;
            margin: 0;
        }
        .metric-subtle {
            color: #61758b;
            font-size: 0.9rem;
        }
        .dashboard-stage {
            border-radius: 18px;
            padding: 1rem 1.05rem;
            border: 1px solid rgba(15, 23, 42, 0.08);
            background: rgba(255, 255, 255, 0.95);
            height: 170px;
            box-shadow: 0 14px 28px rgba(15, 23, 42, 0.05);
            overflow: hidden;
        }
        .dashboard-stage-shell {
            display: flex;
            flex-direction: column;
            gap: 0.5rem;
            height: 100%;
        }
        .dashboard-stage-main {
            flex: 1 1 auto;
        }
        .dashboard-hero {
            padding: 1.2rem 1.35rem;
            border-radius: 22px;
            border: 1px solid rgba(15, 23, 42, 0.08);
            background:
                radial-gradient(circle at top right, rgba(11, 78, 140, 0.16), transparent 32%),
                linear-gradient(135deg, rgba(255, 255, 255, 0.98) 0%, rgba(245, 249, 253, 0.98) 100%);
            box-shadow: 0 18px 36px rgba(15, 23, 42, 0.07);
            margin-bottom: 1rem;
        }
        .dashboard-kicker {
            text-transform: uppercase;
            letter-spacing: 0.08em;
            font-size: 0.72rem;
            color: #0b4e8c;
            font-weight: 800;
        }
        .dashboard-title {
            margin: 0.3rem 0 0.35rem 0;
            font-size: 1.7rem;
            line-height: 1.1;
            color: #10233b;
            font-weight: 800;
        }
        .dashboard-subtitle {
            color: #5a6b80;
            font-size: 0.96rem;
            line-height: 1.45;
        }
        .dashboard-mini-metric {
            border-radius: 18px;
            padding: 1rem 1.05rem;
            border: 1px solid rgba(15, 23, 42, 0.08);
            background: rgba(255, 255, 255, 0.94);
            min-height: 118px;
        }
        .dashboard-mini-label {
            font-size: 0.75rem;
            letter-spacing: 0.06em;
            text-transform: uppercase;
            color: #6d7d91;
            font-weight: 700;
        }
        .dashboard-mini-value {
            margin-top: 0.45rem;
            color: #10233b;
            font-size: 1.5rem;
            font-weight: 800;
            line-height: 1.1;
        }
        .dashboard-mini-copy {
            color: #607287;
            font-size: 0.88rem;
            margin-top: 0.45rem;
        }
        .dashboard-stage-label {
            font-size: 1rem;
            font-weight: 700;
            color: #10233b;
            margin: 0.35rem 0 0.5rem 0;
        }
        .dashboard-stage-copy {
            color: #5d7086;
            font-size: 0.92rem;
            line-height: 1.35;
            min-height: 72px;
            overflow: hidden;
            display: -webkit-box;
            -webkit-line-clamp: 4;
            -webkit-box-orient: vertical;
        }
        div.dashboard-card-action div.stButton > button,
        div.dashboard-card-action button[kind="secondary"],
        div.dashboard-card-action button[kind="primary"] {
            border-radius: 999px;
            min-height: 34px;
            height: 34px;
            padding: 0 0.8rem;
            border: 1px solid rgba(11, 78, 140, 0.9) !important;
            background: linear-gradient(180deg, #0b4e8c 0%, #164f7b 100%) !important;
            color: #ffffff !important;
            font-size: 0.82rem;
            font-weight: 700;
            box-shadow: 0 10px 20px rgba(11, 78, 140, 0.18) !important;
        }
        div.dashboard-card-action div.stButton > button:hover,
        div.dashboard-card-action button[kind="secondary"]:hover,
        div.dashboard-card-action button[kind="primary"]:hover {
            border-color: rgba(11, 78, 140, 1) !important;
            color: #ffffff !important;
            background: linear-gradient(180deg, #0d579c 0%, #0f466d 100%) !important;
        }
        div.dashboard-card-action div.stButton > button p,
        div.dashboard-card-action button[kind="secondary"] p,
        div.dashboard-card-action button[kind="primary"] p {
            color: #ffffff !important;
        }
        .dashboard-badge {
            display: inline-flex;
            align-items: center;
            gap: 0.35rem;
            border-radius: 999px;
            padding: 0.28rem 0.62rem;
            font-size: 0.78rem;
            font-weight: 700;
        }
        .dashboard-badge.done {
            background: rgba(46, 125, 50, 0.12);
            color: #2e7d32;
        }
        .dashboard-badge.running {
            background: rgba(11, 78, 140, 0.12);
            color: #0b4e8c;
        }
        .dashboard-badge.ready {
            background: rgba(239, 108, 0, 0.12);
            color: #ef6c00;
        }
        .dashboard-badge.blocked,
        .dashboard-badge.pending {
            background: rgba(107, 124, 147, 0.12);
            color: #55687e;
        }
        .dashboard-badge.rejected {
            background: rgba(198, 40, 40, 0.12);
            color: #c62828;
        }
        .dashboard-list-card {
            border-radius: 20px;
            padding: 1rem 1.1rem;
            border: 1px solid rgba(15, 23, 42, 0.08);
            background: rgba(255, 255, 255, 0.94);
            box-shadow: 0 14px 30px rgba(15, 23, 42, 0.05);
            height: 100%;
        }
        .dashboard-list-row {
            display: flex;
            justify-content: space-between;
            gap: 0.75rem;
            padding: 0.55rem 0;
            border-bottom: 1px solid rgba(15, 23, 42, 0.07);
        }
        .dashboard-list-row:last-child {
            border-bottom: none;
            padding-bottom: 0;
        }
        .dashboard-list-key {
            color: #66788d;
            font-size: 0.88rem;
        }
        .dashboard-list-value {
            color: #10233b;
            font-size: 0.92rem;
            font-weight: 700;
            text-align: right;
        }
        .dashboard-event {
            border-left: 4px solid #0b4e8c;
            border-radius: 14px;
            padding: 0.8rem 0.9rem;
            background: rgba(255, 255, 255, 0.9);
            border-top: 1px solid rgba(15, 23, 42, 0.06);
            border-right: 1px solid rgba(15, 23, 42, 0.06);
            border-bottom: 1px solid rgba(15, 23, 42, 0.06);
            margin-bottom: 0.65rem;
        }
        .dashboard-event-meta {
            color: #6a7b90;
            font-size: 0.78rem;
            font-weight: 700;
            margin-bottom: 0.2rem;
            text-transform: uppercase;
            letter-spacing: 0.04em;
        }
        .dashboard-event-copy {
            color: #10233b;
            font-size: 0.92rem;
            line-height: 1.35;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_metric_card(label: str, value: str, subtle: str) -> None:
    st.markdown(
        f"""
        <div class="metric-card">
            <div class="small-label">{label}</div>
            <p class="metric-value">{value}</p>
            <div class="metric-subtle">{subtle}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def to_float(value: Any, default: float = 0.0) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        try:
            return float(value.strip())
        except ValueError:
            return default
    return default


def render_bullets(items: list, empty_message: str) -> None:
    if items:
        for item in items:
            st.write(f"- {item}")
    else:
        st.write(empty_message)


def normalize_for_table(items: Any) -> pd.DataFrame:
    if isinstance(items, list):
        if not items:
            return pd.DataFrame()
        if all(isinstance(item, dict) for item in items):
            return pd.json_normalize(items, sep=".")
        return pd.DataFrame({"value": [str(item) for item in items]})
    if isinstance(items, dict):
        return pd.json_normalize(items, sep=".")
    if items in (None, ""):
        return pd.DataFrame()
    return pd.DataFrame({"value": [str(items)]})


DISPLAY_TOKEN_MAP = {
    "api": "API",
    "apis": "APIs",
    "ai": "AI",
    "brd": "BRD",
    "fe": "FE",
    "github": "GitHub",
    "id": "ID",
    "ids": "IDs",
    "ip": "IP",
    "json": "JSON",
    "pr": "PR",
    "qa": "QA",
    "sql": "SQL",
    "sme": "SME",
    "ui": "UI",
    "url": "URL",
    "urls": "URLs",
    "vb": "VB",
}


def format_table_header(label: Any) -> str:
    text = str(label or "").strip()
    if not text:
        return ""
    parts = re.split(r"[._\s]+", text)
    formatted_parts = []
    for part in parts:
        if not part:
            continue
        formatted_parts.append(DISPLAY_TOKEN_MAP.get(part.lower(), part[:1].upper() + part[1:]))
    return " ".join(formatted_parts)


def _looks_like_technical_text(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    return any(marker in stripped for marker in ("://", "\\", "/", "@")) or bool(re.search(r"\.[A-Za-z0-9]{1,8}$", stripped))


def format_table_value(value: Any) -> Any:
    if isinstance(value, str):
        text = value.strip()
        if not text or _looks_like_technical_text(text):
            return value
        if "_" in text and text.lower() == text:
            return format_table_header(text)
        if text[:1].islower():
            return text[:1].upper() + text[1:]
        return text
    return value


def sanitize_dataframe_for_display(dataframe: pd.DataFrame) -> pd.DataFrame:
    if dataframe.empty:
        return dataframe

    def normalize_cell(value: Any) -> Any:
        if isinstance(value, list):
            return ", ".join(str(item) for item in value)
        if isinstance(value, dict):
            return json.dumps(value, ensure_ascii=False, sort_keys=True)
        return format_table_value(value)

    sanitized = dataframe.map(normalize_cell)
    sanitized.columns = [format_table_header(column) for column in sanitized.columns]
    return sanitized


def make_excel_safe_sheet_name(name: str) -> str:
    sanitized = "".join("_" if ch in '[]:*?/\\' else ch for ch in name).strip()
    return (sanitized or "Sheet")[:31]


def prepare_dataframe_for_excel(items: Any) -> pd.DataFrame:
    dataframe = normalize_for_table(items)
    if dataframe.empty:
        return pd.DataFrame({"message": ["No data available."]})
    return sanitize_dataframe_for_display(dataframe)


def write_excel_sheet(writer, sheet_name: str, items: Any) -> None:
    dataframe = prepare_dataframe_for_excel(items)
    dataframe.to_excel(writer, sheet_name=make_excel_safe_sheet_name(sheet_name), index=False)


def write_sectioned_excel_sheet(writer, sheet_name: str, sections: list[tuple[str, Any]]) -> None:
    safe_sheet_name = make_excel_safe_sheet_name(sheet_name)
    start_row = 0
    for title, items in sections:
        title_df = pd.DataFrame([{title: ""}])
        title_df.to_excel(writer, sheet_name=safe_sheet_name, index=False, header=True, startrow=start_row)
        start_row += len(title_df.index) + 2

        dataframe = prepare_dataframe_for_excel(items)
        dataframe.to_excel(writer, sheet_name=safe_sheet_name, index=False, startrow=start_row)
        start_row += len(dataframe.index) + 3


def build_workflow_excel_export(result: dict) -> bytes | None:
    if openpyxl is None or not result:
        return None

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        summary_rows = [
            {"step": "Legacy Code Reverse", "available": bool(result.get("legacy_code_reverse_spec")), "type": "reverse"},
            {"step": "Legacy SQL Reverse", "available": bool(result.get("legacy_sql_reverse_spec")), "type": "reverse"},
            {"step": "Legacy Collate", "available": bool(result.get("legacy_spec")), "type": "collate"},
            {"step": "Target Code Reverse", "available": bool(result.get("target_code_reverse_spec")), "type": "reverse"},
            {"step": "Target SQL Reverse", "available": bool(result.get("target_sql_reverse_spec")), "type": "reverse"},
            {"step": "Target Collate", "available": bool(result.get("target_spec")), "type": "collate"},
            {"step": "Gap Analysis", "available": bool(result.get("gap_analysis")), "type": "gap"},
            {"step": "Requirements Draft", "available": bool(result.get("requirements_draft")), "type": "requirements"},
            {"step": "Technical Spec Draft", "available": bool(result.get("technical_spec_draft")), "type": "technical_spec"},
            {"step": "Forward Engineering", "available": bool(result.get("forward_engineering_output")), "type": "forward_engineering"},
            {"step": "Final Data Mapping", "available": bool(result.get("data_mapping_result")), "type": "data_mapping"},
            {"step": "Validation", "available": bool(result.get("validation_result")), "type": "validation"},
            {"step": "AI Code Review", "available": bool(result.get("ai_code_review_result")), "type": "review"},
        ]
        pd.DataFrame(summary_rows).to_excel(writer, sheet_name="Run Summary", index=False)

        spec_exports = [
            ("1 Legacy Code", result.get("legacy_code_reverse_spec", {})),
            ("2 Legacy SQL", result.get("legacy_sql_reverse_spec", {})),
            ("3 Legacy Collate", result.get("legacy_spec", {})),
            ("4 Target Code", result.get("target_code_reverse_spec", {})),
            ("5 Target SQL", result.get("target_sql_reverse_spec", {})),
            ("6 Target Collate", result.get("target_spec", {})),
        ]
        spec_sections = [
            ("Summary", lambda spec: {"summary": spec.get("summary", ""), "confidence": spec.get("confidence", {}), "source_files": spec.get("source_files", []), "notes": spec.get("notes", [])}),
            ("Fields", lambda spec: spec.get("fields", [])),
            ("Common Rules", lambda spec: spec.get("business_rules", [])),
            ("Country Rules", lambda spec: spec.get("country_specific_rules", [])),
            ("Validations", lambda spec: spec.get("validations", [])),
            ("Calculations", lambda spec: spec.get("calculations", [])),
            ("UI Components", lambda spec: spec.get("ui_components", [])),
            ("Classes", lambda spec: spec.get("classes", [])),
            ("Methods", lambda spec: spec.get("methods", [])),
            ("Procedures", lambda spec: spec.get("procedures", [])),
            ("Procedure Dependencies", lambda spec: spec.get("procedure_dependencies", [])),
            ("Table Dependencies", lambda spec: spec.get("table_dependencies", [])),
            ("API Endpoints", lambda spec: spec.get("api_endpoints", [])),
        ]
        for prefix, spec in spec_exports:
            if not spec:
                continue
            sections = [(section_name, extractor(spec)) for section_name, extractor in spec_sections]
            write_sectioned_excel_sheet(writer, prefix, sections)

        gap_analysis = result.get("gap_analysis", {})
        if gap_analysis:
            write_sectioned_excel_sheet(
                writer,
                "7 Gap Analysis",
                [
                    ("Missing Features", gap_analysis.get("missing_features", [])),
                    ("Incorrect Implementations", gap_analysis.get("incorrect_implementations", [])),
                    ("Compliance Gaps", gap_analysis.get("compliance_gaps", [])),
                    ("Risks", gap_analysis.get("risks", [])),
                    ("Rule Comparison", format_gap_rule_rows(gap_analysis.get("rule_comparison", []))),
                    ("Common Rules Missed", format_gap_rule_rows(gap_analysis.get("common_rules_missed", []))),
                    ("Country Rules Missed", format_gap_rule_rows(gap_analysis.get("country_specific_rules_missed", []))),
                    ("Confidence", gap_analysis.get("confidence", {})),
                ],
            )

        requirements_draft = result.get("requirements_draft", {})
        if requirements_draft:
            write_sectioned_excel_sheet(
                writer,
                "8 Requirements",
                [
                    ("Summary", {
                        "screen_name": requirements_draft.get("screen_name", ""),
                        "business_context": requirements_draft.get("business_context", ""),
                        "assumptions": requirements_draft.get("assumptions", []),
                        "open_questions_for_sme": requirements_draft.get("open_questions_for_sme", []),
                        "review_notes": requirements_draft.get("review_notes", []),
                        "approval": requirements_draft.get("approval", {}),
                    }),
                    ("Functional Requirements", requirements_draft.get("functional_requirements", [])),
                    ("Non Functional Requirements", requirements_draft.get("non_functional_requirements", [])),
                    ("Compliance Requirements", requirements_draft.get("compliance_requirements", [])),
                    ("Data Requirements", requirements_draft.get("data_requirements", [])),
                    ("UI Requirements", requirements_draft.get("ui_requirements", [])),
                    ("API Requirements", requirements_draft.get("api_requirements", [])),
                    ("Migration Requirements", requirements_draft.get("migration_requirements", [])),
                ],
            )

        technical_spec = result.get("technical_spec_draft", {})
        if technical_spec:
            write_sectioned_excel_sheet(
                writer,
                "9 Technical Spec",
                [
                    ("Summary", {
                        "screen_name": technical_spec.get("screen_name", ""),
                        "target_stack": technical_spec.get("target_stack", {}),
                        "assumptions": technical_spec.get("assumptions", []),
                        "open_questions_for_architect": technical_spec.get("open_questions_for_architect", []),
                        "review_notes": technical_spec.get("review_notes", []),
                        "approval": technical_spec.get("approval", {}),
                    }),
                    ("UI Design", technical_spec.get("ui_design", [])),
                    ("API Design", technical_spec.get("api_design", [])),
                    ("Service Design", technical_spec.get("service_design", [])),
                    ("Data Design", technical_spec.get("data_design", [])),
                    ("Rule Configuration Design", technical_spec.get("rule_configuration_design", [])),
                    ("Validation Design", technical_spec.get("validation_design", [])),
                    ("Security And Compliance Design", technical_spec.get("security_and_compliance_design", [])),
                    ("Integration Design", technical_spec.get("integration_design", [])),
                ],
            )

        forward_engineering = result.get("forward_engineering_output", {})
        if forward_engineering:
            write_sectioned_excel_sheet(
                writer,
                "10 Forward Eng",
                [
                    ("Angular Files", forward_engineering.get("angular_files", [])),
                    ("Angular Test Files", forward_engineering.get("angular_test_files", [])),
                    ("Node.js Files", forward_engineering.get("nodejs_files", [])),
                    ("Node.js Test Files", forward_engineering.get("nodejs_test_files", [])),
                    ("PostgreSQL Files", forward_engineering.get("postgres_files", [])),
                    ("Test Cases", forward_engineering.get("test_cases", [])),
                    ("Notes", {
                        "generation_notes": forward_engineering.get("generation_notes", []),
                        "traceability_summary": forward_engineering.get("traceability_summary", []),
                    }),
                ],
            )

        data_mapping_result = result.get("data_mapping_result", {})
        if data_mapping_result:
            write_sectioned_excel_sheet(
                writer,
                "11 Final Data Mapping",
                [
                    ("Summary", data_mapping_result.get("mapping_analysis", {}).get("summary", {})),
                    ("Field Mappings", format_data_mapping_field_rows(data_mapping_result.get("mapping_analysis", {}).get("field_mappings", []))),
                    ("Reconciliation Approach", data_mapping_result.get("mapping_analysis", {}).get("reconciliation_approach", [])),
                    ("Reference Data Mappings", data_mapping_result.get("mapping_analysis", {}).get("reference_data_mappings", [])),
                    ("Unmapped Legacy Fields", data_mapping_result.get("mapping_analysis", {}).get("unmapped_legacy_fields", [])),
                    ("New Final State Fields", data_mapping_result.get("mapping_analysis", {}).get("new_final_state_fields", [])),
                    ("Transformations", data_mapping_result.get("mapping_analysis", {}).get("transformations", [])),
                    ("Data Quality Risks", data_mapping_result.get("mapping_analysis", {}).get("data_quality_risks", [])),
                    ("Migration Notes", data_mapping_result.get("mapping_analysis", {}).get("migration_notes", [])),
                    ("Confidence", data_mapping_result.get("mapping_analysis", {}).get("confidence", {})),
                ],
            )

        ai_code_review_result = result.get("ai_code_review_result", {})
        if ai_code_review_result:
            write_sectioned_excel_sheet(
                writer,
                "12 AI Code Review",
                [
                    ("Summary", ai_code_review_result.get("summary", {})),
                    ("Findings", ai_code_review_result.get("findings", [])),
                    ("Strengths", ai_code_review_result.get("strengths", [])),
                    ("PR Review Comment", {"review_comment_markdown": ai_code_review_result.get("review_comment_markdown", "")}),
                ],
            )

        validation_result = result.get("validation_result", {})
        if validation_result:
            write_sectioned_excel_sheet(
                writer,
                "13 Validation",
                [
                    ("Summary", validation_result.get("summary", {})),
                    ("Differences", validation_result.get("differences", [])),
                    ("Suggestions", validation_result.get("suggestions", [])),
                    ("Confidence", validation_result.get("confidence", {})),
                ],
            )

        if result.get("logs"):
            write_excel_sheet(writer, "Execution Logs", result.get("logs", []))

    return output.getvalue()


def build_nested_rows(items: list[dict], nested_key: str, parent_keys: list[str], child_prefix: str = "") -> list[dict]:
    rows: list[dict] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        nested_items = item.get(nested_key, [])
        if not isinstance(nested_items, list):
            continue
        for nested_item in nested_items:
            row = {key: item.get(key) for key in parent_keys}
            if isinstance(nested_item, dict):
                for key, value in nested_item.items():
                    row[f"{child_prefix}{key}" if child_prefix else key] = value
            else:
                row[f"{child_prefix}value" if child_prefix else "value"] = nested_item
            rows.append(row)
    return rows


def render_table(items: Any, empty_message: str = "No data available.") -> None:
    dataframe = normalize_for_table(items)
    if dataframe.empty:
        st.info(empty_message)
        return
    dataframe = sanitize_dataframe_for_display(dataframe)
    st.dataframe(dataframe, use_container_width=True, hide_index=True)


def render_notes(items: list[Any]) -> None:
    if not items:
        st.info("No notes available.")
        return
    for item in items:
        st.markdown(
            f"""
            <div class="panel-card" style="margin-bottom: 0.7rem;">
                <div>{item}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def create_run_trace_dir() -> Path:
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    run_dir = settings.outputs_dir / "runs" / timestamp
    suffix = 1
    while run_dir.exists():
        suffix += 1
        run_dir = settings.outputs_dir / "runs" / f"{timestamp}_{suffix}"
    run_dir.mkdir(parents=True, exist_ok=True)
    return run_dir


def write_json_trace(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def write_run_trace_snapshot(run_dir: Path, state: dict, status: str) -> None:
    logs = state.get("logs", [])
    write_json_trace(run_dir / "workflow_log.json", {"status": status, "logs": logs})
    write_json_trace(run_dir / "latest_state.json", {"status": status, "state": state})


def finalize_run_trace(run_dir: Path, state: dict) -> None:
    write_run_trace_snapshot(run_dir, state, "completed")
    write_json_trace(run_dir / "final_state.json", state)


def write_error_trace(run_dir: Path, error_message: str, state: dict) -> None:
    write_run_trace_snapshot(run_dir, state, "failed")
    write_json_trace(
        run_dir / "error.json",
        {
            "error": error_message,
            "logs": state.get("logs", []),
            "state": state,
        },
    )


def render_live_execution_status(log_placeholder, logs: list[dict]) -> None:
    with log_placeholder.container():
        st.markdown("**Live Execution Status**")
        if not logs:
            st.info("Waiting for the first workflow update...")
            return
        recent_logs = logs[-8:]
        for entry in recent_logs:
            st.markdown(
                f"- `{entry.get('agent', 'system')}`: {entry.get('message', '')}"
            )


def repo_root() -> Path:
    return Path(__file__).resolve().parent


def parse_github_repo_full_name(remote_url: str) -> str:
    cleaned = remote_url.strip()
    if cleaned.endswith(".git"):
        cleaned = cleaned[:-4]
    https_match = re.search(r"github\.com[:/](?P<owner>[^/]+)/(?P<repo>[^/]+)$", cleaned)
    if https_match:
        return f"{https_match.group('owner')}/{https_match.group('repo')}"
    return ""


def slugify_git_ref(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return slug or "generated-target"


def get_github_token() -> str:
    token = os.environ.get("GITHUB_TOKEN", "").strip()
    # Be forgiving when the token is pasted with surrounding quotes in a .env file or shell.
    if len(token) >= 2 and token[0] == token[-1] and token[0] in {"'", '"'}:
        token = token[1:-1].strip()
    return token


def github_headers(extra: dict[str, str] | None = None) -> dict[str, str]:
    token = get_github_token()
    headers = {
        "Accept": "application/vnd.github+json",
        "User-Agent": "ai-migration-platform",
    }
    if token:
        lowered = token.lower()
        if lowered.startswith("bearer ") or lowered.startswith("token "):
            headers["Authorization"] = token
        elif token.startswith(("ghp_", "github_pat_", "gho_", "ghu_", "ghs_", "ghr_")):
            headers["Authorization"] = f"token {token}"
        else:
            headers["Authorization"] = f"Bearer {token}"
    if extra:
        headers.update(extra)
    return headers


def github_api_json(path: str, method: str = "GET", payload: dict | None = None) -> dict:
    url = path if path.startswith("http") else f"https://api.github.com{path}"
    data = None
    if payload is not None:
        data = json.dumps(payload).encode("utf-8")
    request = urllib_request.Request(url, data=data, method=method, headers=github_headers({"Content-Type": "application/json"}))
    try:
        with urllib_request.urlopen(request) as response:
            raw = response.read().decode("utf-8")
            return json.loads(raw) if raw else {}
    except urllib_error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")
        try:
            message = json.loads(detail).get("message", detail)
        except json.JSONDecodeError:
            message = detail or exc.reason
        if exc.code == 401:
            message = (
                f"{message}. Check GITHUB_TOKEN: it may be expired, pasted with quotes, or missing the required repository permissions."
            )
        raise RuntimeError(f"GitHub API error ({exc.code}): {message}") from exc
    except urllib_error.URLError as exc:
        raise RuntimeError(f"Unable to reach GitHub API: {exc.reason}") from exc


def normalize_repo_path(path: str) -> str:
    return path.strip().strip("/").replace("\\", "/")


def github_fetch_repo_code_snapshot(repo_full_name: str, branch: str, code_root: str, destination: str) -> str:
    if not get_github_token():
        raise RuntimeError("GITHUB_TOKEN is not configured in the environment.")
    normalized_root = normalize_repo_path(code_root)
    branch_payload = github_api_json(f"/repos/{repo_full_name}/branches/{urllib_parse.quote(branch, safe='')}")
    tree_sha = branch_payload.get("commit", {}).get("commit", {}).get("tree", {}).get("sha", "")
    if not tree_sha:
        raise RuntimeError(f"Unable to resolve tree SHA for {repo_full_name}@{branch}.")
    tree_payload = github_api_json(f"/repos/{repo_full_name}/git/trees/{tree_sha}?recursive=1")
    destination_path = Path(destination)
    if destination_path.exists():
        shutil.rmtree(destination_path)
    destination_path.mkdir(parents=True, exist_ok=True)

    for item in tree_payload.get("tree", []):
        if item.get("type") != "blob":
            continue
        repo_path = str(item.get("path", "")).replace("\\", "/")
        if normalized_root and not repo_path.startswith(f"{normalized_root}/") and repo_path != normalized_root:
            continue
        if repo_path.lower().endswith(".sql"):
            continue
        relative_path = repo_path[len(normalized_root):].lstrip("/") if normalized_root else repo_path
        if not relative_path:
            continue
        blob_payload = github_api_json(f"/repos/{repo_full_name}/git/blobs/{item.get('sha', '')}")
        encoded = blob_payload.get("content", "")
        encoding = blob_payload.get("encoding", "base64")
        if encoding != "base64":
            continue
        content = base64.b64decode(encoded.encode("utf-8"))
        file_path = destination_path / Path(relative_path)
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_bytes(content)
    return str(destination_path)


def build_publish_defaults(result: dict, target_branch: str) -> dict:
    screen_name = str(result.get("requirements_draft", {}).get("screen_name", "")).strip()
    branch_name = f"AUTOMATED_FE_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}"
    commit_message = f"Add generated target candidate for {screen_name or 'migration flow'}"
    pr_title = f"[codex] Forward engineering output for {screen_name or 'migration flow'}"
    return {
        "branch_name": branch_name,
        "base_branch": target_branch,
        "commit_message": commit_message,
        "pr_title": pr_title,
    }


def build_forward_engineering_pr_body(result: dict, generated_target: dict, comparison_items: list[dict]) -> str:
    forward_output = result.get("forward_engineering_output", {}) or {}
    added_count = sum(1 for item in comparison_items if item.get("status") == "Added")
    modified_count = sum(1 for item in comparison_items if item.get("status") == "Modified")
    unchanged_count = sum(1 for item in comparison_items if item.get("status") == "Unchanged")
    test_asset_count = len(forward_output.get("angular_test_files", [])) + len(forward_output.get("nodejs_test_files", []))
    traceability = forward_output.get("traceability_summary", [])
    notes = forward_output.get("generation_notes", [])
    generated_root = generated_target.get("generated_root", "")
    body_lines = [
        "## What Changed",
        "Publishes the forward-engineered target candidate generated by the modernization workflow.",
        "",
        "## Generated Output",
        f"- Generated root: `{generated_root}`",
        f"- Added files: {added_count}",
        f"- Modified files: {modified_count}",
        f"- Unchanged files: {unchanged_count}",
        f"- Business test cases: {len(forward_output.get('test_cases', []))}",
        f"- Automated unit test files: {test_asset_count}",
        "",
        "## Why",
        "Captures the AI-generated target implementation in a draft PR so it can be reviewed before merge.",
        "",
        "## Traceability",
    ]
    if traceability:
        body_lines.extend(f"- {item}" for item in traceability[:8])
    else:
        body_lines.append("- Traceability summary not available.")
    body_lines.extend(["", "## Validation"])
    if notes:
        body_lines.extend(f"- {item}" for item in notes[:6])
    else:
        body_lines.append("- Forward engineering generation completed in the app.")
    validation_summary = (result.get("validation_result") or {}).get("summary", {})
    mapping_summary = ((result.get("data_mapping_result") or {}).get("mapping_analysis") or {}).get("summary", {})
    body_lines.extend(["", "## AI-Assisted Code Review Summary"])
    body_lines.append(f"- Added files: {added_count}; modified files: {modified_count}; unchanged files: {unchanged_count}.")
    if validation_summary:
        body_lines.append(
            f"- Validation status: {validation_summary.get('overall_status', 'not_run')} with {validation_summary.get('difference_count', 0)} differences and sync score {to_float(validation_summary.get('sync_score', 0.0)):.0%}."
        )
    else:
        body_lines.append("- Validation has not been run yet in the app.")
    if mapping_summary:
        body_lines.append(
            f"- Data mapping coverage: {to_float(mapping_summary.get('mapping_coverage', 0.0)):.0%}; unmapped columns: {mapping_summary.get('unmapped_columns_count', 0)}."
        )
    else:
        body_lines.append("- Final data mapping has not been run yet in the app.")
    body_lines.append("- SQL artifacts were intentionally excluded from this PR; only code and automated test files are included.")
    return "\n".join(body_lines)


def build_ai_code_review_input(comparison_items: list[dict]) -> list[dict]:
    review_items: list[dict] = []
    for item in comparison_items:
        if item.get("status") not in {"Added", "Modified"}:
            continue
        if item.get("group") == "postgres_files":
            continue
        review_items.append(
            {
                "file_name": item.get("file_name", ""),
                "generated_relative_path": item.get("generated_relative_path", ""),
                "status": item.get("status", ""),
                "group": item.get("group", ""),
                "purpose": item.get("purpose", ""),
                "related_requirement_ids": item.get("related_requirement_ids", []),
                "diff_text": str(item.get("diff_text", ""))[:6000],
                "content_preview": str(item.get("content", ""))[:4000],
            }
        )
    return review_items


def github_post_pr_review_comment(repo_full_name: str, pr_number: int, review_body: str) -> dict:
    return github_api_json(
        f"/repos/{repo_full_name}/pulls/{pr_number}/reviews",
        method="POST",
        payload={"body": review_body, "event": "COMMENT"},
    )


def github_find_open_pr(repo_full_name: str, branch_name: str, base_branch: str) -> dict:
    owner = repo_full_name.split("/", 1)[0]
    pr_items = github_api_json(
        f"/repos/{repo_full_name}/pulls?state=open&head={urllib_parse.quote(f'{owner}:{branch_name}', safe=':')}&base={urllib_parse.quote(base_branch, safe='')}"
    )
    if isinstance(pr_items, list) and pr_items:
        return pr_items[0]
    return {}


def resolve_publish_config(result: dict, target_branch: str) -> dict:
    defaults = build_publish_defaults(result, target_branch)
    persisted_publish = st.session_state.get("github_publish_result") or result.get("github_publish_result") or {}
    return {
        "branch_name": (
            persisted_publish.get("branch_name")
            or st.session_state.get("github_branch_name")
            or ""
        ).strip(),
        "base_branch": (
            persisted_publish.get("base_branch")
            or st.session_state.get("github_base_branch")
            or defaults["base_branch"]
        ).strip(),
        "commit_message": (
            persisted_publish.get("commit_message")
            or st.session_state.get("github_commit_message")
            or defaults["commit_message"]
        ).strip(),
        "pr_title": (
            persisted_publish.get("pr_title")
            or st.session_state.get("github_pr_title")
            or defaults["pr_title"]
        ).strip(),
    }


def publish_generated_target_pr(
    generated_target: dict,
    result: dict,
    comparison_items: list[dict],
    target_repo_full_name: str,
    target_code_root: str,
    branch_name: str,
    base_branch: str,
    commit_message: str,
    pr_title: str,
) -> dict:
    generated_root = generated_target.get("generated_root", "")
    if not generated_root:
        raise RuntimeError("No generated target candidate is available to publish.")
    if not get_github_token():
        raise RuntimeError("GITHUB_TOKEN is not configured in the environment.")
    branch_payload = github_api_json(f"/repos/{target_repo_full_name}/branches/{urllib_parse.quote(base_branch, safe='')}")
    base_sha = branch_payload.get("commit", {}).get("sha", "")
    if not base_sha:
        raise RuntimeError(f"Unable to resolve base branch `{base_branch}` in {target_repo_full_name}.")
    try:
        github_api_json(
            f"/repos/{target_repo_full_name}/git/refs",
            method="POST",
            payload={"ref": f"refs/heads/{branch_name}", "sha": base_sha},
        )
    except RuntimeError as exc:
        if "Reference already exists" not in str(exc) and "422" not in str(exc):
            raise

    normalized_target_root = normalize_repo_path(target_code_root)
    staged_files: list[str] = []
    commit_sha = ""
    for item in comparison_items:
        if item.get("status") not in {"Added", "Modified"}:
            continue
        if item.get("group") == "postgres_files":
            continue
        relative_path = str(item.get("generated_relative_path", "")).replace("\\", "/")
        if not relative_path or relative_path.lower().endswith(".sql"):
            continue
        repo_path = f"{normalized_target_root}/{relative_path}" if normalized_target_root else relative_path
        content = item.get("content", "")
        existing_sha = None
        try:
            existing_payload = github_api_json(
                f"/repos/{target_repo_full_name}/contents/{urllib_parse.quote(repo_path, safe='/')}?ref={urllib_parse.quote(branch_name, safe='')}"
            )
            existing_sha = existing_payload.get("sha")
        except RuntimeError as exc:
            if "404" not in str(exc):
                raise
        update_payload = {
            "message": commit_message,
            "content": base64.b64encode(str(content).encode("utf-8")).decode("utf-8"),
            "branch": branch_name,
        }
        if existing_sha:
            update_payload["sha"] = existing_sha
        update_result = github_api_json(
            f"/repos/{target_repo_full_name}/contents/{urllib_parse.quote(repo_path, safe='/')}",
            method="PUT",
            payload=update_payload,
        )
        staged_files.append(repo_path)
        commit_sha = update_result.get("commit", {}).get("sha", commit_sha)

    if not staged_files:
        raise RuntimeError("No non-SQL generated code or test files were found to publish.")

    try:
        pr_payload = github_api_json(
            f"/repos/{target_repo_full_name}/pulls",
            method="POST",
            payload={
                "title": pr_title,
                "head": branch_name,
                "base": base_branch,
                "body": build_forward_engineering_pr_body(result, generated_target, comparison_items),
                "draft": False,
            },
        )
    except RuntimeError as exc:
        if "A pull request already exists" not in str(exc) and "422" not in str(exc):
            raise
        pr_payload = github_find_open_pr(target_repo_full_name, branch_name, base_branch)
        if not pr_payload:
            raise
    return {
        "repository_full_name": target_repo_full_name,
        "branch_name": branch_name,
        "base_branch": base_branch,
        "commit_message": commit_message,
        "commit_sha": commit_sha,
        "pr_title": pr_title,
        "pr_number": pr_payload.get("number", 0),
        "pr_url": pr_payload.get("html_url", ""),
        "staged_files": staged_files,
    }


def render_github_pr_publish_panel(
    result: dict,
    generated_target: dict,
    comparison_items: list[dict],
    target_repo_full_name: str,
    target_branch: str,
    target_code_root: str,
) -> None:
    st.markdown("**Raise GitHub PR**")
    if not generated_target.get("generated_root"):
        st.info("Generate forward-engineered target artifacts before raising a GitHub PR.")
        return

    defaults = build_publish_defaults(result, target_branch)
    if not st.session_state.get("github_branch_name"):
        st.session_state["github_branch_name"] = defaults["branch_name"]
    if not st.session_state.get("github_commit_message"):
        st.session_state["github_commit_message"] = defaults["commit_message"]
    if not st.session_state.get("github_pr_title"):
        st.session_state["github_pr_title"] = defaults["pr_title"]

    info_col, action_col = st.columns([3, 2])
    with info_col:
        st.caption(f"Repository: {target_repo_full_name}")
        if not get_github_token():
            st.warning("GITHUB_TOKEN is not configured. Set it in the environment before raising a PR.")
    with action_col:
        st.caption(f"Generated folder: `{generated_target.get('generated_root', '')}`")

    base_col, branch_col = st.columns(2)
    with base_col:
        st.text_input("Base branch", key="github_base_branch")
    with branch_col:
        st.text_input("Publish branch", key="github_branch_name")

    st.text_input("Commit message", key="github_commit_message")
    st.text_input("PR title", key="github_pr_title")

    if st.session_state.get("github_publish_error"):
        st.error(st.session_state["github_publish_error"])
    publish_result = st.session_state.get("github_publish_result") or result.get("github_publish_result")
    if publish_result:
        st.success("PR created successfully.")
        publish_rows = [
            {"metric": "Repository", "value": publish_result.get("repository_full_name", "")},
            {"metric": "Branch", "value": publish_result.get("branch_name", "")},
            {"metric": "Base", "value": publish_result.get("base_branch", "")},
            {"metric": "Commit", "value": publish_result.get("commit_sha", "")[:12]},
            {"metric": "PR URL", "value": publish_result.get("pr_url", "") or "Created"},
        ]
        render_table(publish_rows, "No PR details available.")

    publish_disabled = not bool(get_github_token())
    if st.button("Raise GitHub PR", use_container_width=True, disabled=publish_disabled):
        try:
            st.session_state["github_publish_error"] = None
            st.session_state["github_publish_result"] = None
            st.session_state["ai_code_review_result"] = None
            st.session_state["ai_code_review_error"] = None
            publish_result = publish_generated_target_pr(
                generated_target=generated_target,
                result=result,
                comparison_items=comparison_items,
                target_repo_full_name=target_repo_full_name,
                target_code_root=target_code_root,
                branch_name=st.session_state.get("github_branch_name", defaults["branch_name"]).strip(),
                base_branch=st.session_state.get("github_base_branch", defaults["base_branch"]).strip(),
                commit_message=st.session_state.get("github_commit_message", defaults["commit_message"]).strip(),
                pr_title=st.session_state.get("github_pr_title", defaults["pr_title"]).strip(),
            )
            st.session_state["github_publish_result"] = publish_result
            st.session_state["analysis_result"]["github_publish_result"] = publish_result
            st.session_state["pending_main_section"] = "AI Code Review"
            st.rerun()
        except Exception as exc:
            st.session_state["github_publish_error"] = str(exc)
            st.rerun()


def summarize_forward_engineering_artifacts(output: dict | None) -> int:
    if not isinstance(output, dict):
        return 0
    return (
        len(output.get("angular_files", []))
        + len(output.get("angular_test_files", []))
        + len(output.get("nodejs_files", []))
        + len(output.get("nodejs_test_files", []))
        + len(output.get("postgres_files", []))
    )


def format_elapsed_from_logs(logs: list[dict]) -> str:
    if len(logs) < 2:
        return "Just started"
    timestamps: list[datetime] = []
    for entry in logs:
        raw = entry.get("timestamp")
        if not isinstance(raw, str):
            continue
        try:
            timestamps.append(datetime.fromisoformat(raw.replace("Z", "+00:00")))
        except ValueError:
            continue
    if len(timestamps) < 2:
        return "In progress"
    elapsed = max((max(timestamps) - min(timestamps)).total_seconds(), 0)
    minutes, seconds = divmod(int(elapsed), 60)
    hours, minutes = divmod(minutes, 60)
    if hours:
        return f"{hours}h {minutes}m"
    if minutes:
        return f"{minutes}m {seconds}s"
    return f"{seconds}s"


def _approval_status_for_dashboard(state: dict, approval_name: str) -> str:
    if approval_name == "requirements":
        document = state.get("approved_requirements") or state.get("requirements_draft")
        return get_approval_status(document, PENDING_SME_APPROVAL)
    document = state.get("approved_technical_spec") or state.get("technical_spec_draft")
    return get_approval_status(document, PENDING_ARCHITECT_APPROVAL)


def _latest_log_entry(logs: list[dict], agent_name: str) -> dict:
    for entry in reversed(logs):
        if entry.get("agent") == agent_name:
            return entry
    return {}


def _is_active_running_log(entry: dict) -> bool:
    message = str(entry.get("message", "")).strip().lower()
    if not message:
        return False
    terminal_markers = ("complete", "completed", "skipped", "reused", "cache hit", "approved", "rejected")
    return not any(marker in message for marker in terminal_markers)


def _approval_detail(state: dict, approval_name: str) -> str:
    if approval_name == "requirements":
        document = state.get("approved_requirements") or state.get("requirements_draft") or {}
        role = "SME"
    else:
        document = state.get("approved_technical_spec") or state.get("technical_spec_draft") or {}
        role = "Architect"
    approval = document.get("approval", {}) if isinstance(document, dict) else {}
    status = approval.get("status", "")
    approved_on = approval.get("approved_on", "")
    comments = approval.get("review_comments", "")
    if status and "APPROVED" in str(status):
        return f"Approved by {role}{f' on {approved_on}' if approved_on else ''}."
    if status and "REJECTED" in str(status):
        return comments or f"Rejected by {role}."
    if document:
        return f"Awaiting {role} approval."
    return f"Pending {role} review."


def _dashboard_stage_status(stage: dict, state: dict, live: bool = False) -> tuple[str, str]:
    logs = state.get("logs", [])
    latest_entry = _latest_log_entry(logs, stage["key"])
    stage_key = stage["key"]
    if stage.get("approval") == "requirements":
        approval_status = _approval_status_for_dashboard(state, "requirements")
        if approval_status == SME_APPROVED:
            return "done", "Complete"
        if state.get("requirements_draft"):
            return "pending", "Pending"
        if live and _is_active_running_log(_latest_log_entry(logs, "requirements")):
            return "running", "Running"
        return "pending", "Pending"

    if stage.get("approval") == "technical_spec":
        approval_status = _approval_status_for_dashboard(state, "technical_spec")
        if approval_status == ARCHITECT_APPROVED:
            return "done", "Complete"
        if state.get("technical_spec_draft"):
            return "pending", "Pending"
        if live and _is_active_running_log(_latest_log_entry(logs, "technical_spec")):
            return "running", "Running"
        return "pending", "Pending"

    if all(state.get(key) for key in stage.get("result_keys", [])):
        return "done", "Complete"
    if stage_key == "forward_engineering_proof":
        if state.get("forward_comparison_items"):
            return "done", "Complete"
        if state.get("forward_engineering_output"):
            return "pending", "Pending"
        return "pending", "Pending"
    if stage_key == "technical_spec" and _approval_status_for_dashboard(state, "requirements") != SME_APPROVED:
        return "pending", "Pending"
    if stage_key == "forward_engineering":
        if _approval_status_for_dashboard(state, "requirements") != SME_APPROVED:
            return "pending", "Pending"
        if _approval_status_for_dashboard(state, "technical_spec") != ARCHITECT_APPROVED:
            return "pending", "Pending"
    if stage_key in {"data_mapping", "validation"} and not state.get("forward_engineering_output"):
        return "pending", "Pending"
    if live and _is_active_running_log(latest_entry):
        return "running", "Running"
    return "pending", "Pending"


def build_dashboard_rows(state: dict, live: bool = False) -> list[dict]:
    logs = state.get("logs", [])
    rows: list[dict] = []
    for stage in DASHBOARD_STAGES:
        status_key, status_label = _dashboard_stage_status(stage, state, live=live)
        detail = _latest_stage_message(logs, stage["key"])
        if stage.get("approval"):
            detail = _approval_detail(state, stage["approval"])
        elif not detail and status_key == "done":
            detail = "Completed."
        elif not detail and status_key == "pending":
            detail = "Pending."
        rows.append(
            {
                "stage": stage["label"],
                "status_key": status_key,
                "status": status_label,
                "detail": detail,
                "detail_section": stage.get("detail_section", ""),
                "approval": stage.get("approval", ""),
            }
        )
    return rows


def _latest_stage_message(logs: list[dict], stage_key: str) -> str:
    for entry in reversed(logs):
        if entry.get("agent") == stage_key:
            return entry.get("message", "")
    return ""


def render_dashboard_metric(label: str, value: str, copy: str) -> None:
    st.markdown(
        f"""
        <div class="dashboard-mini-metric">
            <div class="dashboard-mini-label">{label}</div>
            <div class="dashboard-mini-value">{value}</div>
            <div class="dashboard-mini-copy">{copy}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_dashboard_snapshot(snapshot_rows: list[dict]) -> None:
    rows_markup = "".join(
        f"""
        <div class="dashboard-list-row">
            <div class="dashboard-list-key">{row['metric']}</div>
            <div class="dashboard-list-value">{row['value']}</div>
        </div>
        """
        for row in snapshot_rows
    )
    st.markdown(
        f"""
        <div class="dashboard-list-card">
            <div class="section-title">Run Snapshot</div>
            {rows_markup}
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_dashboard_events(logs: list[dict], limit: int = 6) -> None:
    if not logs:
        st.info("The dashboard will populate as soon as execution begins.")
        return

    for entry in logs[-limit:]:
        timestamp = entry.get("timestamp", "")
        try:
            timestamp = datetime.fromisoformat(timestamp.replace("Z", "+00:00")).strftime("%H:%M:%S")
        except ValueError:
            pass
        st.markdown(
            f"""
            <div class="dashboard-event">
                <div class="dashboard-event-meta">{timestamp} | {entry.get('agent', 'system')}</div>
                <div class="dashboard-event-copy">{entry.get('message', '')}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def go_to_dashboard_approval_section(section_name: str) -> None:
    st.session_state["active_main_section"] = section_name
    st.rerun()


def run_final_data_mapping_from_workspace(
    result: dict,
    generated_target: dict,
    legacy_data_dictionary_folder: str,
    target_data_dictionary_folder: str,
) -> None:
    st.session_state["data_mapping_error"] = None
    st.session_state["ai_code_review_result"] = None
    st.session_state["ai_code_review_error"] = None
    generated_sql_folder = str(Path(generated_target.get("generated_root", "")) / "sql")
    with st.spinner("Running final data mapping analysis on the generated target..."):
        data_mapping_result = data_mapping_agent.run(
            legacy_spec=result.get("legacy_spec", {}),
            target_spec=result.get("target_spec", {}),
            generated_code_folder=generated_target.get("generated_root", ""),
            generated_sql_folder=generated_sql_folder,
            legacy_data_dictionary_folder=legacy_data_dictionary_folder,
            target_data_dictionary_folder=target_data_dictionary_folder,
            logger=lambda _a, _m: None,
        )
    st.session_state["data_mapping_result"] = data_mapping_result
    st.session_state["analysis_result"]["data_mapping_result"] = data_mapping_result
    st.session_state["pending_main_section"] = "Final Data Mapping"
    st.rerun()


def run_validation_from_workspace(
    result: dict,
    requirements_draft: dict,
    technical_spec_draft: dict,
    generated_target: dict,
) -> None:
    st.session_state["validation_error"] = None
    st.session_state["ai_code_review_result"] = None
    st.session_state["ai_code_review_error"] = None
    with st.spinner("Running validation on generated target artifacts..."):
        validation_result = validation_agent.run(
            legacy_spec=result.get("legacy_spec", {}),
            requirements=requirements_draft,
            technical_spec=technical_spec_draft,
            generated_files=generated_target.get("generated_files", []),
            logger=lambda _a, _m: None,
        )
    st.session_state["validation_result"] = validation_result
    st.session_state["analysis_result"]["validation_result"] = validation_result
    st.session_state["pending_main_section"] = "Validation"
    st.rerun()


def render_dashboard_stage_cards(rows: list[dict], live: bool = False) -> None:
    current_state = st.session_state.get("analysis_result", {})
    stage_rows = [rows[i : i + 4] for i in range(0, len(rows), 4)]
    for group_index, row_group in enumerate(stage_rows):
        cols = st.columns(4)
        for index, row in enumerate(row_group):
            with cols[index]:
                detail = row["detail"] or "No updates yet."
                st.markdown(
                    f"""
                    <div class="dashboard-stage">
                        <div class="dashboard-stage-shell">
                            <div class="dashboard-stage-main">
                                <span class="dashboard-badge {row['status_key']}">{row['status']}</span>
                                <div class="dashboard-stage-label">{row['stage']}</div>
                                <div class="dashboard-stage-copy">{detail}</div>
                            </div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                if not live and row.get("detail_section"):
                    st.markdown('<div class="dashboard-card-action">', unsafe_allow_html=True)
                    button_key = f"dashboard_open_{group_index}_{index}_{row['stage']}"
                    if st.button("↗ Details", key=button_key, use_container_width=True):
                        st.session_state["active_main_section"] = row["detail_section"]
                        st.rerun()
                    st.markdown("</div>", unsafe_allow_html=True)
                if (
                    not live
                    and row.get("approval") == "requirements"
                    and row["status_key"] == "pending"
                    and current_state.get("requirements_draft")
                ):
                    st.markdown('<div class="dashboard-card-action">', unsafe_allow_html=True)
                    if st.button("✓ Review", key=f"dashboard_approve_req_{group_index}_{index}", use_container_width=True):
                        go_to_dashboard_approval_section("Requirements Draft")
                    st.markdown("</div>", unsafe_allow_html=True)
                if (
                    not live
                    and row.get("approval") == "technical_spec"
                    and row["status_key"] == "pending"
                    and current_state.get("technical_spec_draft")
                ):
                    st.markdown('<div class="dashboard-card-action">', unsafe_allow_html=True)
                    if st.button("✓ Review", key=f"dashboard_approve_tech_{group_index}_{index}", use_container_width=True):
                        go_to_dashboard_approval_section("Technical Specification Draft")
                    st.markdown("</div>", unsafe_allow_html=True)
        for empty_index in range(len(row_group), 4):
            cols[empty_index].empty()


def dashboard_mapping_analysis(state: dict) -> dict:
    data_mapping_result = state.get("data_mapping_result") or {}
    if not isinstance(data_mapping_result, dict):
        return {}
    mapping_analysis = data_mapping_result.get("mapping_analysis", {})
    return mapping_analysis if isinstance(mapping_analysis, dict) else {}


def dashboard_validation_summary(state: dict) -> dict:
    validation_result = state.get("validation_result") or {}
    if not isinstance(validation_result, dict):
        return {}
    summary = validation_result.get("summary", {})
    return summary if isinstance(summary, dict) else {}


def compute_dashboard_readiness(state: dict) -> dict:
    gap_analysis = state.get("gap_analysis", {}) or {}
    mapping_analysis = dashboard_mapping_analysis(state)
    mapping_summary = mapping_analysis.get("summary", {}) if isinstance(mapping_analysis, dict) else {}
    validation_summary = dashboard_validation_summary(state)
    requirements_status = _approval_status_for_dashboard(state, "requirements")
    technical_status = _approval_status_for_dashboard(state, "technical_spec")

    business_readiness = 100
    if gap_analysis.get("missing_features"):
        business_readiness -= min(40, len(gap_analysis.get("missing_features", [])) * 6)
    if requirements_status != SME_APPROVED:
        business_readiness -= 20

    technical_readiness = 100
    if technical_status != ARCHITECT_APPROVED:
        technical_readiness -= 30
    if not state.get("forward_engineering_output"):
        technical_readiness -= 25
    technical_readiness -= min(25, int(validation_summary.get("difference_count", 0)) * 4)

    data_readiness = 100
    data_readiness -= min(40, int(mapping_summary.get("unmapped_columns_count", 0)) * 6)
    data_readiness -= min(25, len(mapping_analysis.get("data_quality_risks", [])) * 5)
    if not state.get("data_mapping_result"):
        data_readiness -= 25

    compliance_readiness = 100
    compliance_readiness -= min(45, len(gap_analysis.get("compliance_gaps", [])) * 10)
    compliance_readiness -= min(20, len(gap_analysis.get("risks", [])) * 4)

    testing_readiness = 100
    if not state.get("validation_result"):
        testing_readiness -= 35
    testing_readiness -= min(35, int(validation_summary.get("difference_count", 0)) * 5)
    testing_readiness += min(10, len((state.get("forward_engineering_output") or {}).get("test_cases", [])))

    areas = {
        "Business Readiness": max(0, min(100, business_readiness)),
        "Technical Readiness": max(0, min(100, technical_readiness)),
        "Data Readiness": max(0, min(100, data_readiness)),
        "Compliance Readiness": max(0, min(100, compliance_readiness)),
        "Testing Readiness": max(0, min(100, testing_readiness)),
    }
    overall_score = round(sum(areas.values()) / len(areas)) if areas else 0
    if overall_score >= 80:
        health = "Green"
        wave_status = "Release candidate"
    elif overall_score >= 60:
        health = "Amber"
        wave_status = "Wave candidate with follow-ups"
    else:
        health = "Red"
        wave_status = "Not ready for release"
    return {"areas": areas, "overall_score": overall_score, "health": health, "wave_status": wave_status}


def build_blockers_and_risks(state: dict) -> list[dict]:
    blockers: list[dict] = []
    gap_analysis = state.get("gap_analysis", {}) or {}
    if _approval_status_for_dashboard(state, "requirements") != SME_APPROVED:
        blockers.append({"type": "Approval", "severity": "High", "item": "BRD pending SME approval"})
    if _approval_status_for_dashboard(state, "technical_spec") != ARCHITECT_APPROVED:
        blockers.append({"type": "Approval", "severity": "High", "item": "Technical spec pending architect approval"})
    for item in gap_analysis.get("compliance_gaps", [])[:3]:
        blockers.append({"type": "Compliance", "severity": "High", "item": str(item)})
    for risk in gap_analysis.get("risks", [])[:4]:
        if isinstance(risk, dict):
            blockers.append(
                {
                    "type": "Risk",
                    "severity": risk.get("risk_level", "Medium"),
                    "item": risk.get("risk_name", risk.get("description", "Open risk")),
                }
            )
        else:
            blockers.append({"type": "Risk", "severity": "Medium", "item": str(risk)})
    return blockers


def build_delivery_progress_rows(state: dict) -> list[dict]:
    mapping_analysis = dashboard_mapping_analysis(state)
    mapping_summary = mapping_analysis.get("summary", {}) if isinstance(mapping_analysis, dict) else {}
    validation_summary = dashboard_validation_summary(state)
    rows = []
    for row in build_dashboard_rows(state):
        progress = 100 if row["status_key"] == "done" else 75 if row["status_key"] == "ready" else 50 if row["status_key"] == "running" else 0
        rows.append({"step": row["stage"], "status": row["status"], "progress_%": progress})
    rows.extend(
        [
            {"step": "Gap Closure Status", "status": f"{len((state.get('gap_analysis') or {}).get('missing_features', []))} gaps open", "progress_%": max(0, 100 - len((state.get("gap_analysis") or {}).get("missing_features", [])) * 10)},
            {"step": "Data Mapping Coverage", "status": f"{to_float(mapping_summary.get('mapping_coverage', 0.0)):.0%}", "progress_%": round(to_float(mapping_summary.get("mapping_coverage", 0.0)) * 100)},
            {"step": "Validation Sync", "status": validation_summary.get("overall_status", "not_run"), "progress_%": round(to_float(validation_summary.get("sync_score", 0.0)) * 100)},
        ]
    )
    return rows


def build_country_readiness_rows(state: dict) -> list[dict]:
    gap_analysis = state.get("gap_analysis", {}) or {}
    legacy_spec = state.get("legacy_spec", {}) or {}
    country_names = {
        str(item.get("country", "")).strip()
        for item in legacy_spec.get("country_specific_rules", [])
        if isinstance(item, dict) and str(item.get("country", "")).strip()
    }
    missed_by_country: dict[str, int] = {}
    for item in gap_analysis.get("country_specific_rules_missed", []):
        if not isinstance(item, dict):
            continue
        country = str(item.get("country", "")).strip() or "Unspecified"
        country_names.add(country)
        missed_by_country[country] = missed_by_country.get(country, 0) + 1
    rows = []
    for country in sorted(country_names):
        missed = missed_by_country.get(country, 0)
        readiness = max(0, 100 - missed * 25)
        rows.append({"country": country, "rules_missed": missed, "readiness_%": readiness, "status": "Ready" if readiness >= 75 else "Needs attention"})
    return rows


def build_module_readiness_rows(state: dict) -> list[dict]:
    output = state.get("forward_engineering_output", {}) or {}
    validation_summary = dashboard_validation_summary(state)
    rows = [
        {"module": "Frontend", "artifacts": len(output.get("angular_files", [])), "status": "Ready" if output.get("angular_files") else "Pending"},
        {"module": "Frontend Tests", "artifacts": len(output.get("angular_test_files", [])), "status": "Ready" if output.get("angular_test_files") else "Pending"},
        {"module": "Backend", "artifacts": len(output.get("nodejs_files", [])), "status": "Ready" if output.get("nodejs_files") else "Pending"},
        {"module": "Backend Tests", "artifacts": len(output.get("nodejs_test_files", [])), "status": "Ready" if output.get("nodejs_test_files") else "Pending"},
        {"module": "Database", "artifacts": len(output.get("postgres_files", [])), "status": "Ready" if output.get("postgres_files") else "Pending"},
        {"module": "Validation", "artifacts": int(validation_summary.get("difference_count", 0)), "status": validation_summary.get("overall_status", "not_run")},
    ]
    return rows


def build_file_generation_rows(state: dict) -> list[dict]:
    items = state.get("forward_comparison_items", []) or []
    rows = []
    for item in items:
        rows.append(
            {
                "file_name": item.get("file_name", ""),
                "status": item.get("status", ""),
                "module": str(item.get("group", "")).replace("_files", "").upper(),
                "path": item.get("generated_relative_path", ""),
                "requirements": ", ".join(str(req) for req in item.get("related_requirement_ids", [])),
            }
        )
    return rows


def render_migration_dashboard(state: dict, live: bool = False) -> None:
    logs = state.get("logs", [])
    rows = build_dashboard_rows(state, live=live)
    done_count = sum(1 for row in rows if row["status_key"] == "done")
    action_count = sum(1 for row in rows if row.get("approval") and row["status_key"] == "pending")
    progress_ratio = done_count / len(rows) if rows else 0.0
    forward_output = state.get("forward_engineering_output", {})
    validation_result = state.get("validation_result") or {}
    mapping_analysis = dashboard_mapping_analysis(state)
    mapping_summary = mapping_analysis.get("summary", {}) if isinstance(mapping_analysis, dict) else {}
    validation_summary = dashboard_validation_summary(state)
    readiness = compute_dashboard_readiness(state)
    blockers = build_blockers_and_risks(state)
    generated_count = summarize_forward_engineering_artifacts(forward_output)
    dashboard_title = "Live Migration Command Center" if live else "Migration Command Center"
    dashboard_copy = (
        "Streaming workflow telemetry, approval checkpoints, and generated asset readiness."
        if live
        else "A consolidated view of execution progress, review gates, and migration delivery readiness."
    )
    st.markdown(
        f"""
        <div class="dashboard-hero">
            <div class="dashboard-kicker">Migration Dashboard</div>
            <div class="dashboard-title">{dashboard_title}</div>
            <div class="dashboard-subtitle">{dashboard_copy}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    has_readiness_evidence = any(
        state.get(key)
        for key in [
            "logs",
            "legacy_code_reverse_spec",
            "legacy_sql_reverse_spec",
            "legacy_spec",
            "target_code_reverse_spec",
            "target_sql_reverse_spec",
            "target_spec",
            "gap_analysis",
            "requirements_draft",
            "technical_spec_draft",
            "forward_engineering_output",
            "validation_result",
            "data_mapping_result",
        ]
    )

    progress_text = "Live migration execution" if live else "Migration readiness"
    if has_readiness_evidence:
        st.progress(progress_ratio, text=f"{progress_text}: {progress_ratio:.0%}")
    else:
        st.caption("Readiness percentages appear after analysis begins.")

    snapshot_rows = [
        {"metric": "Elapsed", "value": format_elapsed_from_logs(logs)},
        {"metric": "Log Events", "value": len(logs)},
        {"metric": "Open Blockers", "value": len(blockers)},
        {"metric": "Generated Assets", "value": generated_count},
        {"metric": "Validation Differences", "value": validation_summary.get("difference_count", 0)},
        {"metric": "Mapped Columns", "value": mapping_summary.get("mapped_columns_count", 0)},
    ]
    executive_tab, delivery_tab, data_tab, validation_tab, drilldown_tab = st.tabs(
        ["Executive Summary", "Delivery Progress", "Data Migration", "Validation and Defects", "Drilldowns"]
    )

    with executive_tab:
        top_left, top_right = st.columns([2, 3])
        with top_left:
            render_dashboard_snapshot(snapshot_rows)
        with top_right:
            st.markdown("**Executive Phase View**")
            render_dashboard_stage_cards(rows, live=live)
            st.markdown("**Recent Events**")
            render_dashboard_events(logs, limit=5 if live else 7)

    with delivery_tab:
        progress_left, progress_right = st.columns([3, 2])
        with progress_left:
            st.markdown("**Delivery Progress**")
            render_table(build_delivery_progress_rows(state), "No delivery progress available.")
        with progress_right:
            delivery_rows = [
                {"metric": "BRD Approval", "value": _approval_status_for_dashboard(state, "requirements")},
                {"metric": "Tech Spec Approval", "value": _approval_status_for_dashboard(state, "technical_spec")},
                {"metric": "Forward Engineering", "value": "Complete" if state.get("forward_engineering_output") else "Pending"},
                {"metric": "Data Mapping", "value": "Complete" if state.get("data_mapping_result") else "Pending"},
                {"metric": "Validation", "value": validation_summary.get("overall_status", "not_run")},
                {"metric": "Test Cases", "value": len(forward_output.get("test_cases", []))},
            ]
            render_dashboard_snapshot(delivery_rows)

    with data_tab:
        data_left, data_right = st.columns([3, 2])
        with data_left:
            st.markdown("**Data Migration Overview**")
            data_rows = [
                {"metric": "Mapped Columns", "value": mapping_summary.get("mapped_columns_count", 0)},
                {"metric": "Unmapped Columns", "value": mapping_summary.get("unmapped_columns_count", 0)},
                {"metric": "Transformations Defined", "value": mapping_summary.get("transformation_count", 0)},
                {"metric": "Reconciliation Rules Defined", "value": mapping_summary.get("reconciliation_rule_count", 0)},
                {"metric": "Coverage", "value": f"{to_float(mapping_summary.get('mapping_coverage', 0.0)):.0%}"},
            ]
            render_table(data_rows, "No data migration metrics available.")
            st.markdown("**Reconciliation Progress**")
            render_table(mapping_analysis.get("reconciliation_approach", []), "No reconciliation approach generated.")
        with data_right:
            st.markdown("**Migration Risks**")
            render_table(mapping_analysis.get("data_quality_risks", []), "No data quality risks detected.")
            st.markdown("**Unmapped Fields**")
            render_table(mapping_analysis.get("unmapped_legacy_fields", []), "No unmapped legacy fields detected.")

    with validation_tab:
        validation_left, validation_right = st.columns([2, 3])
        with validation_left:
            validation_rows = [
                {"metric": "Overall Status", "value": validation_summary.get("overall_status", "not_run")},
                {"metric": "Sync Score", "value": f"{to_float(validation_summary.get('sync_score', 0.0)):.0%}"},
                {"metric": "Total Issues", "value": validation_summary.get("difference_count", 0)},
                {"metric": "Open Suggestions", "value": validation_summary.get("suggestion_count", 0)},
                {"metric": "Generated Inconsistencies", "value": len(validation_result.get("differences", []))},
            ]
            render_dashboard_snapshot(validation_rows)
        with validation_right:
            st.markdown("**Validation and Defects**")
            render_table(validation_result.get("differences", []), "No validation defects detected.")
            st.markdown("**Recommended Actions**")
            render_table(validation_result.get("suggestions", []), "No corrective actions suggested.")

    with drilldown_tab:
        drill_tabs = st.tabs(["By Country", "By Module", "By Generated File", "By Workflow Step"])
        with drill_tabs[0]:
            render_table(build_country_readiness_rows(state), "No country-specific readiness indicators available.")
        with drill_tabs[1]:
            render_table(build_module_readiness_rows(state), "No module-level readiness indicators available.")
        with drill_tabs[2]:
            render_table(build_file_generation_rows(state), "No generated file status available.")
        with drill_tabs[3]:
            workflow_rows = []
            for row in rows:
                workflow_rows.append({"workflow_step": row["stage"], "status": row["status"], "latest_update": row["detail"] or ""})
            render_table(workflow_rows, "No workflow-step detail available.")


def deep_copy_document(document: dict | None) -> dict:
    if not isinstance(document, dict):
        return {}
    return json.loads(json.dumps(document))


def get_approval_status(document: dict | None, default_status: str) -> str:
    if not isinstance(document, dict):
        return default_status
    approval = document.get("approval", {})
    if not isinstance(approval, dict):
        return default_status
    status = approval.get("status", default_status)
    return status if isinstance(status, str) and status else default_status


def update_document_approval(document: dict | None, status: str, comments: str, reviewer_label: str) -> dict:
    updated = deep_copy_document(document)
    updated.setdefault("approval", {})
    updated["approval"]["status"] = status
    updated["approval"]["required_reviewer_role"] = reviewer_label
    updated["approval"]["approved_by"] = reviewer_label if "APPROVED" in status else ""
    updated["approval"]["approved_on"] = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    updated["approval"]["review_comments"] = comments
    return updated


def init_approval_state() -> None:
    st.session_state.setdefault("approved_requirements", None)
    st.session_state.setdefault("approved_technical_spec", None)
    st.session_state.setdefault("sme_comments", "")
    st.session_state.setdefault("architect_comments", "")
    st.session_state.setdefault("auto_run_analysis", False)
    st.session_state.setdefault("auto_run_stage", "")
    st.session_state.setdefault("validation_result", None)
    st.session_state.setdefault("validation_error", None)
    st.session_state.setdefault("data_mapping_result", None)
    st.session_state.setdefault("data_mapping_error", None)
    st.session_state.setdefault("ai_code_review_result", None)
    st.session_state.setdefault("ai_code_review_error", None)
    st.session_state.setdefault("github_publish_result", None)
    st.session_state.setdefault("github_publish_error", None)
    st.session_state.setdefault("github_base_branch", "main")
    st.session_state.setdefault("github_branch_name", "")
    st.session_state.setdefault("github_commit_message", "")
    st.session_state.setdefault("github_pr_title", "")
    st.session_state.setdefault("pending_main_section", "")
    st.session_state.setdefault("show_analysis_workspace", False)


def sync_approval_state_from_result(result: dict) -> None:
    requirements_draft = result.get("requirements_draft", {})
    technical_spec_draft = result.get("technical_spec_draft", {})

    approved_requirements = st.session_state.get("approved_requirements")
    if approved_requirements is None and get_approval_status(requirements_draft, PENDING_SME_APPROVAL) == SME_APPROVED:
        st.session_state["approved_requirements"] = deep_copy_document(requirements_draft)

    approved_technical_spec = st.session_state.get("approved_technical_spec")
    if approved_technical_spec is None and get_approval_status(technical_spec_draft, PENDING_ARCHITECT_APPROVAL) == ARCHITECT_APPROVED:
        st.session_state["approved_technical_spec"] = deep_copy_document(technical_spec_draft)


def render_approval_summary(label: str, status: str, comments: str) -> None:
    st.markdown(f"**{label} Status:** `{status}`")
    if comments.strip():
        st.caption(f"Review comments: {comments.strip()}")


def _stringify_value(value: Any) -> str:
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    if isinstance(value, list):
        return ", ".join(str(item) for item in value)
    return str(value)


def _add_doc_section(document, title: str, value: Any) -> None:
    document.add_heading(title, level=1)
    if isinstance(value, list):
        if not value:
            document.add_paragraph("None")
            return
        for item in value:
            if isinstance(item, dict):
                paragraph = document.add_paragraph(style="List Bullet")
                first = True
                for key, item_value in item.items():
                    text = f"{key}: {_stringify_value(item_value)}"
                    if first:
                        paragraph.add_run(text)
                        first = False
                    else:
                        document.add_paragraph(text, style="List Bullet 2")
            else:
                document.add_paragraph(_stringify_value(item), style="List Bullet")
        return

    if isinstance(value, dict):
        if not value:
            document.add_paragraph("None")
            return
        for key, item_value in value.items():
            document.add_paragraph(f"{key}: {_stringify_value(item_value)}", style="List Bullet")
        return

    document.add_paragraph(_stringify_value(value) if value not in (None, "") else "None")


def _set_cell_text(cell, text: Any, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold


def _shade_table_header(table, fill: str = "D9EAF7") -> None:
    if OxmlElement is None or qn is None or not table.rows:
        return
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def _format_document_base(document, title: str, subtitle: str) -> None:
    if Pt is not None:
        for section in document.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        styles = document.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(10)
        styles["Title"].font.name = "Calibri"
        styles["Title"].font.size = Pt(20)

    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH is not None else None
    title_run = title_para.add_run(title)
    title_run.bold = True
    if Pt is not None:
        title_run.font.size = Pt(20)

    subtitle_para = document.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH is not None else None
    subtitle_run = subtitle_para.add_run(subtitle)
    if Pt is not None:
        subtitle_run.font.size = Pt(11)

    document.add_paragraph("")


def _add_metadata_table(document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_cell_text(table.rows[0].cells[0], "Document Attribute", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Value", bold=True)
    _shade_table_header(table)
    for label, value in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], label, bold=True)
        _set_cell_text(cells[1], _stringify_value(value) if value not in (None, "") else "Not provided")
    document.add_paragraph("")


def _add_paragraph_section(document, heading: str, body: Any) -> None:
    document.add_heading(heading, level=1)
    if isinstance(body, list):
        if not body:
            document.add_paragraph("None.")
            return
        for item in body:
            document.add_paragraph(_stringify_value(item), style="List Bullet")
        return
    if isinstance(body, dict):
        if not body:
            document.add_paragraph("None.")
            return
        for key, value in body.items():
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{key}: ").bold = True
            paragraph.add_run(_stringify_value(value))
        return
    text = str(body).strip() if body not in (None, "") else ""
    document.add_paragraph(text or "None.")


def _add_key_value_table_section(document, heading: str, items: list[dict], field_labels: list[tuple[str, str]]) -> None:
    document.add_heading(heading, level=1)
    if not items:
        document.add_paragraph("No entries recorded.")
        return

    for index, item in enumerate(items, start=1):
        document.add_heading(f"{heading[:-1] if heading.endswith('s') else heading} {index}", level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        _set_cell_text(table.rows[0].cells[0], "Attribute", bold=True)
        _set_cell_text(table.rows[0].cells[1], "Detail", bold=True)
        _shade_table_header(table)
        for label, key in field_labels:
            value = item.get(key, "") if isinstance(item, dict) else ""
            cells = table.add_row().cells
            _set_cell_text(cells[0], label, bold=True)
            if isinstance(value, list):
                _set_cell_text(cells[1], "\n".join(str(entry) for entry in value) if value else "None")
            else:
                _set_cell_text(cells[1], _stringify_value(value) if value not in (None, "") else "None")
        document.add_paragraph("")


def _add_matrix_table_section(document, heading: str, items: list[dict], columns: list[tuple[str, str]]) -> None:
    document.add_heading(heading, level=1)
    if not items:
        document.add_paragraph("No entries recorded.")
        return

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, (label, _key) in enumerate(columns):
        _set_cell_text(table.rows[0].cells[idx], label, bold=True)
    _shade_table_header(table)

    for item in items:
        row_cells = table.add_row().cells
        for idx, (_label, key) in enumerate(columns):
            value = item.get(key, "") if isinstance(item, dict) else ""
            if isinstance(value, list):
                value = "\n".join(str(entry) for entry in value)
            _set_cell_text(row_cells[idx], _stringify_value(value) if value not in (None, "") else "None")
    document.add_paragraph("")


def build_requirements_docx(requirements_draft: dict) -> bytes | None:
    if not requirements_draft or Document is None:
        return None

    document = Document()
    _format_document_base(document, "Business Requirements Document", "Modernization Program Working Draft")
    _add_metadata_table(
        document,
        [
            ("Document Type", "Business Requirements Document"),
            ("Status", requirements_draft.get("approval", {}).get("status", "PENDING_SME_APPROVAL")),
            ("Required Reviewer", requirements_draft.get("approval", {}).get("required_reviewer_role", "SME")),
            ("Screen / Process", requirements_draft.get("screen_name", "Quote Generation")),
            ("Generated On", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")),
        ],
    )

    _add_paragraph_section(document, "1. Business Context", requirements_draft.get("business_context", ""))
    _add_matrix_table_section(
        document,
        "2. Functional Requirements",
        requirements_draft.get("functional_requirements", []),
        [
            ("ID", "id"),
            ("Title", "title"),
            ("Description", "description"),
            ("Priority", "priority"),
            ("Rationale", "rationale"),
            ("Source Gap", "source_gap"),
            ("Acceptance Criteria", "acceptance_criteria"),
        ],
    )
    _add_matrix_table_section(
        document,
        "3. Non-Functional Requirements",
        requirements_draft.get("non_functional_requirements", []),
        [("ID", "id"), ("Title", "title"), ("Description", "description"), ("Priority", "priority")],
    )
    _add_matrix_table_section(
        document,
        "4. Compliance Requirements",
        requirements_draft.get("compliance_requirements", []),
        [("ID", "id"), ("Title", "title"), ("Description", "description"), ("Region", "region"), ("Priority", "priority")],
    )
    _add_matrix_table_section(
        document,
        "5. Data Requirements",
        requirements_draft.get("data_requirements", []),
        [("ID", "id"), ("Entity", "entity"), ("Requirement", "requirement")],
    )
    _add_matrix_table_section(
        document,
        "6. UI Requirements",
        requirements_draft.get("ui_requirements", []),
        [("ID", "id"), ("Field / Component", "field_or_component"), ("Requirement", "requirement")],
    )
    _add_matrix_table_section(
        document,
        "7. API Requirements",
        requirements_draft.get("api_requirements", []),
        [("ID", "id"), ("API Name", "api_name"), ("Requirement", "requirement")],
    )
    _add_matrix_table_section(
        document,
        "8. Migration Requirements",
        requirements_draft.get("migration_requirements", []),
        [("ID", "id"), ("Requirement", "requirement")],
    )
    _add_paragraph_section(document, "9. Assumptions", requirements_draft.get("assumptions", []))
    _add_paragraph_section(document, "10. Open Questions For SME", requirements_draft.get("open_questions_for_sme", []))
    _add_paragraph_section(document, "11. Review Notes", requirements_draft.get("review_notes", []))
    _add_key_value_table_section(
        document,
        "12. Approval Summary",
        [requirements_draft.get("approval", {})],
        [
            ("Status", "status"),
            ("Required Reviewer Role", "required_reviewer_role"),
            ("Approved By", "approved_by"),
            ("Approved On", "approved_on"),
            ("Review Comments", "review_comments"),
        ],
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_technical_spec_docx(technical_spec_draft: dict) -> bytes | None:
    if not technical_spec_draft or Document is None:
        return None

    document = Document()
    _format_document_base(document, "Technical Specification", "Modernization Program Working Draft")
    _add_metadata_table(
        document,
        [
            ("Document Type", "Technical Specification"),
            ("Status", technical_spec_draft.get("approval", {}).get("status", "PENDING_ARCHITECT_APPROVAL")),
            ("Required Reviewer", technical_spec_draft.get("approval", {}).get("required_reviewer_role", "Architect")),
            ("Screen / Process", technical_spec_draft.get("screen_name", "Quote Generation")),
            ("Generated On", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")),
        ],
    )

    _add_key_value_table_section(
        document,
        "1. Target Stack",
        [technical_spec_draft.get("target_stack", {})],
        [("Frontend", "frontend"), ("Backend", "backend"), ("Database", "database")],
    )
    _add_matrix_table_section(
        document,
        "2. UI Design",
        technical_spec_draft.get("ui_design", []),
        [("Component", "component"), ("Responsibility", "responsibility"), ("Related Requirement IDs", "related_requirement_ids")],
    )
    _add_matrix_table_section(
        document,
        "3. API Design",
        technical_spec_draft.get("api_design", []),
        [
            ("Name", "name"),
            ("Method", "method"),
            ("Path", "path"),
            ("Request Fields", "request_fields"),
            ("Response Fields", "response_fields"),
            ("Related Requirement IDs", "related_requirement_ids"),
        ],
    )
    _add_matrix_table_section(
        document,
        "4. Service Design",
        technical_spec_draft.get("service_design", []),
        [
            ("Service Name", "service_name"),
            ("Responsibility", "responsibility"),
            ("Business Rules Supported", "business_rules_supported"),
            ("Related Requirement IDs", "related_requirement_ids"),
        ],
    )
    _add_matrix_table_section(
        document,
        "5. Data Design",
        technical_spec_draft.get("data_design", []),
        [("Entity", "entity"), ("Fields", "fields"), ("Relationships", "relationships"), ("Related Requirement IDs", "related_requirement_ids")],
    )
    _add_matrix_table_section(
        document,
        "6. Rule Configuration Design",
        technical_spec_draft.get("rule_configuration_design", []),
        [("Rule Area", "rule_area"), ("Approach", "approach"), ("Details", "details"), ("Related Requirement IDs", "related_requirement_ids")],
    )
    _add_matrix_table_section(
        document,
        "7. Validation Design",
        technical_spec_draft.get("validation_design", []),
        [("Validation Name", "validation_name"), ("Logic", "logic"), ("Layer", "layer"), ("Related Requirement IDs", "related_requirement_ids")],
    )
    _add_matrix_table_section(
        document,
        "8. Security And Compliance Design",
        technical_spec_draft.get("security_and_compliance_design", []),
        [("Area", "area"), ("Design Decision", "design_decision"), ("Related Requirement IDs", "related_requirement_ids")],
    )
    _add_matrix_table_section(
        document,
        "9. Integration Design",
        technical_spec_draft.get("integration_design", []),
        [("Integration Point", "integration_point"), ("Details", "details")],
    )
    _add_paragraph_section(document, "10. Assumptions", technical_spec_draft.get("assumptions", []))
    _add_paragraph_section(document, "11. Open Questions For Architect", technical_spec_draft.get("open_questions_for_architect", []))
    _add_paragraph_section(document, "12. Review Notes", technical_spec_draft.get("review_notes", []))
    _add_key_value_table_section(
        document,
        "13. Approval Summary",
        [technical_spec_draft.get("approval", {})],
        [
            ("Status", "status"),
            ("Required Reviewer Role", "required_reviewer_role"),
            ("Approved By", "approved_by"),
            ("Approved On", "approved_on"),
            ("Review Comments", "review_comments"),
        ],
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_word_download_button(document_label: str, file_name: str, document_bytes: bytes | None, button_key: str) -> None:
    if Document is None:
        st.info("Install `python-docx` to enable Word document downloads.")
        return
    if not document_bytes:
        st.info(f"{document_label} download will be available once the document is generated.")
        return
    st.download_button(
        label=f"Download {document_label}",
        data=document_bytes,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=button_key,
        use_container_width=True,
    )


def filter_country_specific_requirements(functional_requirements: list[dict]) -> list[dict]:
    return [
        item
        for item in functional_requirements
        if isinstance(item, dict) and str(item.get("id", "")).startswith("FR-COUNTRY")
    ]


def filter_country_specific_technical_design(technical_spec_draft: dict) -> dict[str, list[dict]]:
    country_requirement_ids = {
        str(item.get("id", ""))
        for item in filter_country_specific_requirements(technical_spec_draft.get("source_requirements", []))
        if isinstance(item, dict)
    }
    if not country_requirement_ids:
        for section_name in (
            "ui_design",
            "api_design",
            "service_design",
            "data_design",
            "rule_configuration_design",
            "validation_design",
            "security_and_compliance_design",
        ):
            for item in technical_spec_draft.get(section_name, []):
                if not isinstance(item, dict):
                    continue
                related = item.get("related_requirement_ids", [])
                if isinstance(related, list):
                    for requirement_id in related:
                        requirement_id_text = str(requirement_id)
                        if requirement_id_text.startswith("FR-COUNTRY"):
                            country_requirement_ids.add(requirement_id_text)

    def _filter(section_name: str) -> list[dict]:
        results: list[dict] = []
        for item in technical_spec_draft.get(section_name, []):
            if not isinstance(item, dict):
                continue
            related = item.get("related_requirement_ids", [])
            if isinstance(related, list) and any(str(req_id) in country_requirement_ids for req_id in related):
                results.append(item)
        return results

    return {
        "ui_design": _filter("ui_design"),
        "api_design": _filter("api_design"),
        "service_design": _filter("service_design"),
        "data_design": _filter("data_design"),
        "rule_configuration_design": _filter("rule_configuration_design"),
        "validation_design": _filter("validation_design"),
        "security_and_compliance_design": _filter("security_and_compliance_design"),
    }


def render_requirements_document(requirements_draft: dict) -> None:
    if not requirements_draft:
        st.info("Requirements draft will appear after analysis runs.")
        return

    render_word_download_button(
        "BRD (.docx)",
        "business_requirements_document.docx",
        build_requirements_docx(requirements_draft),
        "download_brd_docx",
    )
    st.markdown("**Overview**")
    st.write(requirements_draft.get("business_context", "No business context available."))

    top1, top2, top3 = st.columns(3)
    top1.metric("Functional", len(requirements_draft.get("functional_requirements", [])))
    top2.metric("Open SME Questions", len(requirements_draft.get("open_questions_for_sme", [])))
    top3.metric("Compliance", len(requirements_draft.get("compliance_requirements", [])))
    summary_tab, functional_tab, supporting_tab, approval_tab = st.tabs(
        ["Summary", "Functional Requirements", "Supporting Requirements", "Approval"]
    )

    with summary_tab:
        st.markdown("**Screen Name**")
        st.write(requirements_draft.get("screen_name", ""))
        st.markdown("**Assumptions**")
        render_table(requirements_draft.get("assumptions", []), "No assumptions recorded.")
        st.markdown("**Open Questions For SME**")
        render_table(requirements_draft.get("open_questions_for_sme", []), "No open questions recorded.")
        st.markdown("**Review Notes**")
        render_table(requirements_draft.get("review_notes", []), "No review notes recorded.")

    with functional_tab:
        render_table(requirements_draft.get("functional_requirements", []), "No functional requirements generated.")

    with supporting_tab:
        st.markdown("**Non-Functional Requirements**")
        render_table(requirements_draft.get("non_functional_requirements", []), "No non-functional requirements generated.")
        st.markdown("**Compliance Requirements**")
        render_table(requirements_draft.get("compliance_requirements", []), "No compliance requirements generated.")
        st.markdown("**Data Requirements**")
        render_table(requirements_draft.get("data_requirements", []), "No data requirements generated.")
        st.markdown("**UI Requirements**")
        render_table(requirements_draft.get("ui_requirements", []), "No UI requirements generated.")
        st.markdown("**API Requirements**")
        render_table(requirements_draft.get("api_requirements", []), "No API requirements generated.")
        st.markdown("**Migration Requirements**")
        render_table(requirements_draft.get("migration_requirements", []), "No migration requirements generated.")

    with approval_tab:
        render_table(requirements_draft.get("approval", {}), "No approval metadata available.")


def render_technical_spec_document(technical_spec_draft: dict) -> None:
    if not technical_spec_draft:
        st.info("Technical specification draft will appear after SME approval and the next analysis run.")
        return

    render_word_download_button(
        "Technical Specification (.docx)",
        "technical_specification.docx",
        build_technical_spec_docx(technical_spec_draft),
        "download_technical_spec_docx",
    )
    st.markdown("**Overview**")
    st.write(f"Target stack: {json.dumps(technical_spec_draft.get('target_stack', {}), ensure_ascii=False)}")

    top1, top2, top3 = st.columns(3)
    top1.metric("UI Design", len(technical_spec_draft.get("ui_design", [])))
    top2.metric("API Design", len(technical_spec_draft.get("api_design", [])))
    top3.metric("Architect Questions", len(technical_spec_draft.get("open_questions_for_architect", [])))

    summary_tab, design_tab, review_tab = st.tabs(["Summary", "Design Details", "Approval"])

    with summary_tab:
        st.markdown("**Screen Name**")
        st.write(technical_spec_draft.get("screen_name", ""))
        st.markdown("**Target Stack**")
        render_table(technical_spec_draft.get("target_stack", {}), "No stack details available.")
        st.markdown("**Assumptions**")
        render_table(technical_spec_draft.get("assumptions", []), "No assumptions recorded.")
        st.markdown("**Open Questions For Architect**")
        render_table(technical_spec_draft.get("open_questions_for_architect", []), "No architect questions recorded.")
        st.markdown("**Review Notes**")
        render_table(technical_spec_draft.get("review_notes", []), "No review notes recorded.")

    with design_tab:
        st.markdown("**UI Design**")
        render_table(technical_spec_draft.get("ui_design", []), "No UI design generated.")
        st.markdown("**API Design**")
        render_table(technical_spec_draft.get("api_design", []), "No API design generated.")
        st.markdown("**Service Design**")
        render_table(technical_spec_draft.get("service_design", []), "No service design generated.")
        st.markdown("**Data Design**")
        render_table(technical_spec_draft.get("data_design", []), "No data design generated.")
        st.markdown("**Rule Configuration Design**")
        render_table(technical_spec_draft.get("rule_configuration_design", []), "No rule configuration design generated.")
        st.markdown("**Validation Design**")
        render_table(technical_spec_draft.get("validation_design", []), "No validation design generated.")
        st.markdown("**Security And Compliance Design**")
        render_table(
            technical_spec_draft.get("security_and_compliance_design", []),
            "No security and compliance design generated.",
        )
        st.markdown("**Integration Design**")
        render_table(technical_spec_draft.get("integration_design", []), "No integration design generated.")

    with review_tab:
        render_table(technical_spec_draft.get("approval", {}), "No approval metadata available.")


def render_forward_engineering_document(output: dict) -> None:
    if not output:
        st.info("Forward engineering output will appear after architect approval and the next analysis run.")
        return

    top1, top2, top3, top4 = st.columns(4)
    top1.metric("Angular Files", len(output.get("angular_files", [])))
    top2.metric("Node.js Files", len(output.get("nodejs_files", [])))
    top3.metric("PostgreSQL Files", len(output.get("postgres_files", [])))
    top4.metric("Test Assets", len(output.get("angular_test_files", [])) + len(output.get("nodejs_test_files", [])))

    files_tab, tests_tab, notes_tab = st.tabs(["Generated Files", "Test Cases", "Notes"])

    with files_tab:
        st.markdown("**Angular Files**")
        render_table(output.get("angular_files", []), "No Angular files generated.")
        st.markdown("**Angular Test Files**")
        render_table(output.get("angular_test_files", []), "No Angular test files generated.")
        st.markdown("**Node.js Files**")
        render_table(output.get("nodejs_files", []), "No Node.js files generated.")
        st.markdown("**Node.js Test Files**")
        render_table(output.get("nodejs_test_files", []), "No Node.js test files generated.")
        st.markdown("**PostgreSQL Files**")
        render_table(output.get("postgres_files", []), "No PostgreSQL files generated.")

    with tests_tab:
        st.markdown("**Automated Unit Test Files**")
        automated_test_rows = [
            *output.get("angular_test_files", []),
            *output.get("nodejs_test_files", []),
        ]
        render_table(automated_test_rows, "No automated unit test files generated.")
        st.markdown("**Test Case Catalog**")
        render_table(output.get("test_cases", []), "No test cases generated.")

    with notes_tab:
        st.markdown("**Generation Notes**")
        render_table(output.get("generation_notes", []), "No generation notes available.")
        st.markdown("**Traceability Summary**")
        render_table(output.get("traceability_summary", []), "No traceability summary available.")


def get_generated_target_root() -> Path:
    return settings.outputs_dir / "forward_engineering" / "target_candidate"


def _is_relative_to(path: Path, other: Path) -> bool:
    try:
        path.relative_to(other)
        return True
    except ValueError:
        return False


def _copy_folder_contents(source: Path, destination: Path) -> None:
    if not source.exists() or not source.is_dir():
        return
    destination.mkdir(parents=True, exist_ok=True)
    for item in source.iterdir():
        target = destination / item.name
        if item.is_dir():
            shutil.copytree(item, target, dirs_exist_ok=True)
        else:
            shutil.copy2(item, target)


def seed_generated_target_root(generated_root: Path, target_code_folder: str, target_sql_folder: str) -> None:
    target_code_root = Path(target_code_folder)
    target_sql_root = Path(target_sql_folder)

    if target_code_root.exists() and target_code_root.is_dir():
        _copy_folder_contents(target_code_root, generated_root)

    if (
        target_sql_root.exists()
        and target_sql_root.is_dir()
        and not _is_relative_to(target_sql_root, target_code_root)
    ):
        sql_destination = generated_root / target_sql_root.name
        _copy_folder_contents(target_sql_root, sql_destination)


def build_generated_artifact_path(
    group_name: str,
    file_name: str,
    target_code_folder: str,
    target_sql_folder: str,
) -> Path:
    root = get_generated_target_root()
    safe_name = Path(file_name).name
    search_roots = [target_code_folder] if group_name != "postgres_files" else [target_sql_folder, target_code_folder]
    matched_target_file = find_matching_target_file(safe_name, search_roots)

    if matched_target_file:
        for search_root in search_roots:
            root_path = Path(search_root)
            if root_path.exists() and _is_relative_to(matched_target_file, root_path):
                relative_path = matched_target_file.relative_to(root_path)
                if group_name == "postgres_files" and search_root == target_sql_folder and not _is_relative_to(
                    Path(target_sql_folder), Path(target_code_folder)
                ):
                    return root / Path(target_sql_folder).name / relative_path
                return root / relative_path

    if group_name == "angular_files":
        return root / "frontend" / "src" / "app" / "quote-generation" / safe_name
    if group_name == "angular_test_files":
        return root / "frontend" / "src" / "app" / "quote-generation" / safe_name
    if group_name == "nodejs_files":
        return root / "backend" / "src" / "services" / safe_name
    if group_name == "nodejs_test_files":
        return root / "backend" / "src" / "services" / safe_name
    return root / "sql" / safe_name


def materialize_forward_engineering_output(output: dict, target_code_folder: str, target_sql_folder: str) -> dict:
    cache_payload = {
        "output": output,
        "target_code_folder": target_code_folder,
        "target_sql_folder": target_sql_folder,
    }
    cache = AgentCache(settings.cache_dir)
    if settings.cache_enabled:
        cached = cache.load("forward_engineering_ui", cache_payload)
        if cached:
            generated_root = Path(cached.get("generated_root", ""))
            generated_files = cached.get("generated_files", [])
            if generated_root.exists() and all(Path(item.get("generated_path", "")).exists() for item in generated_files):
                return cached

    generated_root = get_generated_target_root()
    if generated_root.exists():
        shutil.rmtree(generated_root)
    generated_root.mkdir(parents=True, exist_ok=True)
    seed_generated_target_root(generated_root, target_code_folder, target_sql_folder)

    generated_files: list[dict] = []
    for group_name in ("angular_files", "angular_test_files", "nodejs_files", "nodejs_test_files", "postgres_files"):
        for item in output.get(group_name, []):
            file_name = item.get("file_name", "")
            if not file_name:
                continue
            destination = build_generated_artifact_path(
                group_name,
                file_name,
                target_code_folder,
                target_sql_folder,
            )
            destination.parent.mkdir(parents=True, exist_ok=True)
            relative_generated_path = destination.relative_to(generated_root)
            destination.write_text(item.get("content", ""), encoding="utf-8")
            generated_files.append(
                {
                    "group": group_name,
                    "file_name": file_name,
                    "generated_path": str(destination),
                    "generated_relative_path": str(relative_generated_path),
                    "purpose": item.get("purpose", ""),
                    "related_requirement_ids": item.get("related_requirement_ids", []),
                    "content": item.get("content", ""),
                }
            )

    materialized = {"generated_root": str(generated_root), "generated_files": generated_files}
    if settings.cache_enabled:
        cache.save("forward_engineering_ui", cache_payload, materialized)
    return materialized


def find_matching_target_file(file_name: str, search_roots: list[str]) -> Path | None:
    for root in search_roots:
        root_path = Path(root)
        if not root_path.exists():
            continue
        direct_path = root_path / file_name
        if direct_path.exists():
            return direct_path
        matches = list(root_path.rglob(file_name))
        if matches:
            return matches[0]
    return None


def find_original_file_for_generated_artifact(item: dict, target_code_folder: str, target_sql_folder: str) -> Path | None:
    generated_relative_path = item.get("generated_relative_path", "")
    target_code_root = Path(target_code_folder)
    target_sql_root = Path(target_sql_folder)

    if generated_relative_path:
        relative_path = Path(generated_relative_path)
        candidate_in_code_root = target_code_root / relative_path
        if candidate_in_code_root.exists():
            return candidate_in_code_root

        if relative_path.parts:
            if relative_path.parts[0] == target_sql_root.name:
                sql_relative_path = Path(*relative_path.parts[1:]) if len(relative_path.parts) > 1 else Path()
                candidate_in_sql_root = target_sql_root / sql_relative_path
                if candidate_in_sql_root.exists():
                    return candidate_in_sql_root
            else:
                candidate_in_sql_root = target_sql_root / relative_path
                if candidate_in_sql_root.exists():
                    return candidate_in_sql_root

    return find_matching_target_file(item.get("file_name", ""), [target_code_folder, target_sql_folder])


def build_forward_engineering_comparison(generated_files: list[dict], target_code_folder: str, target_sql_folder: str) -> list[dict]:
    comparisons: list[dict] = []

    for item in generated_files:
        file_name = item.get("file_name", "")
        generated_content = item.get("content", "")
        original_path = find_original_file_for_generated_artifact(item, target_code_folder, target_sql_folder)
        original_content = original_path.read_text(encoding="utf-8") if original_path else ""
        if not original_path:
            status = "Added"
        elif original_content == generated_content:
            status = "Unchanged"
        else:
            status = "Modified"

        diff_text = "\n".join(
            difflib.unified_diff(
                original_content.splitlines(),
                generated_content.splitlines(),
                fromfile=f"original/{file_name}",
                tofile=f"generated/{file_name}",
                lineterm="",
            )
        )

        comparisons.append(
            {
                **item,
                "status": status,
                "original_path": str(original_path) if original_path else "",
                "original_content": original_content,
                "diff_text": diff_text,
            }
        )

    return comparisons


def render_forward_engineering_summary(output: dict, comparison_items: list[dict], generated_root: str) -> None:
    if not output:
        st.info("Forward engineering output will appear after architect approval and the next analysis run.")
        return

    added_count = sum(1 for item in comparison_items if item.get("status") == "Added")
    modified_count = sum(1 for item in comparison_items if item.get("status") == "Modified")
    unchanged_count = sum(1 for item in comparison_items if item.get("status") == "Unchanged")

    top1, top2, top3, top4 = st.columns(4)
    top1.metric("Added Files", added_count)
    top2.metric("Modified Files", modified_count)
    top3.metric("Unchanged Files", unchanged_count)
    top4.metric("Total Generated", str(len(comparison_items)))

    st.caption(f"Generated target candidate path: {generated_root}")

    summary_tab, files_tab, tests_tab, notes_tab = st.tabs(["Change Summary", "Generated Files", "Test Cases", "Notes"])

    with summary_tab:
        summary_rows = [
            {
                "status": item.get("status", ""),
                "file_name": item.get("file_name", ""),
                "generated_relative_path": item.get("generated_relative_path", ""),
                "purpose": item.get("purpose", ""),
                "generated_path": item.get("generated_path", ""),
                "original_path": item.get("original_path", ""),
            }
            for item in comparison_items
        ]
        render_table(summary_rows, "No generated files available.")

    with files_tab:
        st.markdown("**Angular Files**")
        render_table(output.get("angular_files", []), "No Angular files generated.")
        st.markdown("**Angular Test Files**")
        render_table(output.get("angular_test_files", []), "No Angular test files generated.")
        st.markdown("**Node.js Files**")
        render_table(output.get("nodejs_files", []), "No Node.js files generated.")
        st.markdown("**Node.js Test Files**")
        render_table(output.get("nodejs_test_files", []), "No Node.js test files generated.")
        st.markdown("**PostgreSQL Files**")
        render_table(output.get("postgres_files", []), "No PostgreSQL files generated.")

    with tests_tab:
        st.markdown("**Automated Unit Test Files**")
        automated_test_rows = [
            *output.get("angular_test_files", []),
            *output.get("nodejs_test_files", []),
        ]
        render_table(automated_test_rows, "No automated unit test files generated.")
        st.markdown("**Test Case Catalog**")
        render_table(output.get("test_cases", []), "No test cases generated.")

    with notes_tab:
        st.markdown("**Generation Notes**")
        render_table(output.get("generation_notes", []), "No generation notes available.")
        st.markdown("**Traceability Summary**")
        render_table(output.get("traceability_summary", []), "No traceability summary available.")


def render_forward_engineering_proof(comparison_items: list[dict]) -> None:
    if not comparison_items:
        st.info("Forward engineering proof will appear after forward engineering generates files.")
        return

    file_options = [item.get("file_name", f"generated_{index}") for index, item in enumerate(comparison_items, start=1)]
    selected_file = st.selectbox("Generated artifact", options=file_options, key="forward_proof_file")
    selected_artifact = next((item for item in comparison_items if item.get("file_name") == selected_file), {})

    generated_content = selected_artifact.get("content", "")
    original_path_value = selected_artifact.get("original_path", "")
    original_path = Path(original_path_value) if original_path_value else None
    original_content = selected_artifact.get("original_content", "")

    original_tab, generated_tab, diff_tab = st.tabs(["Original", "Generated", "Diff"])
    with original_tab:
        if original_path:
            st.caption(f"Matched target file: {original_path}")
            st.code(original_content, language="text")
        else:
            st.info("No matching target file found for this generated artifact.")

    with generated_tab:
        if selected_artifact.get("generated_relative_path"):
            st.caption(f"Generated target path: {selected_artifact.get('generated_relative_path')}")
        st.code(generated_content, language="text")

    with diff_tab:
        if original_path:
            diff_text = selected_artifact.get("diff_text", "")
            st.code(diff_text or "No differences detected.", language="diff")
        else:
            st.info("Diff is unavailable because no original target file match was found.")


def render_validation_result(validation_result: dict) -> None:
    if not validation_result:
        st.info("Run validation after forward engineering to compare generated target artifacts against the source behavior.")
        return

    summary = validation_result.get("summary", {})
    top1, top2, top3 = st.columns(3)
    top1.metric("Sync Score", f"{to_float(summary.get('sync_score', 0.0)):.0%}")
    top2.metric("Differences", str(summary.get("difference_count", 0)))
    top3.metric("Suggestions", str(summary.get("suggestion_count", 0)))

    notes = summary.get("notes", [])
    if notes:
        st.caption(" ".join(str(item) for item in notes))

    st.markdown("**Differences**")
    render_table(validation_result.get("differences", []), "No material differences detected.")

    st.markdown("**Suggestions**")
    render_table(validation_result.get("suggestions", []), "No suggestions generated.")


def render_ai_code_review_result(ai_code_review_result: dict) -> None:
    if not ai_code_review_result:
        st.info("Run AI Code Review after publishing the draft PR to generate and post review findings.")
        return

    summary = ai_code_review_result.get("summary", {})
    top1, top2, top3, top4 = st.columns(4)
    top1.metric("Recommendation", str(summary.get("overall_recommendation", "comment")).replace("_", " ").title())
    top2.metric("Risk Level", str(summary.get("risk_level", "medium")).title())
    top3.metric("Findings", str(summary.get("total_findings", 0)))
    top4.metric("Critical", str(summary.get("critical_findings", 0)))

    notes = summary.get("notes", [])
    if notes:
        st.caption(" ".join(str(item) for item in notes))

    findings_tab, strengths_tab, comment_tab = st.tabs(["Findings", "Strengths", "PR Review Comment"])
    with findings_tab:
        render_table(ai_code_review_result.get("findings", []), "No review findings were generated.")
    with strengths_tab:
        render_table(ai_code_review_result.get("strengths", []), "No strengths were recorded.")
    with comment_tab:
        st.code(ai_code_review_result.get("review_comment_markdown", ""), language="markdown")


def render_data_mapping_result(data_mapping_result: dict) -> None:
    if not data_mapping_result:
        st.info("Run final data mapping after forward engineering proof to compare the final generated design with the legacy and initial target data models.")
        return

    mapping_analysis = data_mapping_result.get("mapping_analysis", {})
    summary = mapping_analysis.get("summary", {})
    top1, top2, top3, top4 = st.columns(4)
    top1.metric("Mapping Coverage", f"{to_float(summary.get('mapping_coverage', 0.0)):.0%}")
    top2.metric("Mapped Columns", str(summary.get("mapped_columns_count", 0)))
    top3.metric("Unmapped Columns", str(summary.get("unmapped_columns_count", 0)))
    top4.metric("Transformations", str(summary.get("transformation_count", 0)))

    findings = summary.get("key_findings", [])
    if findings:
        st.caption(" ".join(str(item) for item in findings))

    mapping_tab, reconciliation_tab, risk_tab = st.tabs(["Column Mapping", "Reconciliation", "Risks"])
    with mapping_tab:
        st.markdown("**Column Mapping**")
        render_table(format_data_mapping_field_rows(mapping_analysis.get("field_mappings", [])), "No field mappings generated.")
        st.markdown("**Transformation Logic**")
        render_table(mapping_analysis.get("transformations", []), "No transformations generated.")
        st.markdown("**Reference Data Mappings**")
        render_table(mapping_analysis.get("reference_data_mappings", []), "No reference data mappings generated.")
        st.markdown("**New Final State Fields**")
        render_table(mapping_analysis.get("new_final_state_fields", []), "No new final-state fields recorded.")
    with reconciliation_tab:
        st.markdown("**Reconciliation Approach**")
        render_table(mapping_analysis.get("reconciliation_approach", []), "No reconciliation approach generated.")
        st.markdown("**Unmapped Legacy Fields**")
        render_table(mapping_analysis.get("unmapped_legacy_fields", []), "No unmapped legacy fields detected.")
    with risk_tab:
        st.markdown("**Data Quality Risks**")
        render_table(mapping_analysis.get("data_quality_risks", []), "No data quality risks detected.")
        st.markdown("**Migration Notes**")
        render_table(mapping_analysis.get("migration_notes", []), "No migration notes generated.")


def format_gap_rule_rows(items: list[dict], prefix: str = "BR") -> list[dict]:
    rows: list[dict] = []
    for index, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        row = dict(item)
        canonical_rule_key = str(row.get("rule_id", "")).strip()
        row["rule_id"] = f"{prefix}-{index:03d}"
        row["Canonical Rule Key"] = canonical_rule_key.replace("*", "_") if canonical_rule_key else ""
        rows.append(row)
    return rows


def format_data_mapping_field_rows(items: list[dict]) -> list[dict]:
    rows: list[dict] = []
    ordered_keys = [
        "final_target_entity",
        "final_target_field",
        "data_mapping_rule",
        "legacy_entity",
        "legacy_field",
        "initial_target_entity",
        "initial_target_field",
        "mapping_type",
        "migration_status",
        "change_category",
        "source_data_type",
        "initial_target_data_type",
        "final_target_data_type",
        "data_source",
        "allowed_values",
        "required_for_migration",
        "notes",
        "evidence",
    ]
    for item in items:
        if not isinstance(item, dict):
            continue
        normalized_item = dict(item)
        normalized_item["data_mapping_rule"] = normalized_item.get("data_mapping_rule") or normalized_item.get("transformation_logic", "")
        row: dict[str, Any] = {}
        for key in ordered_keys:
            row[key] = normalized_item.get(key, "")
        for key, value in normalized_item.items():
            if key not in row and key != "transformation_logic":
                row[key] = value
        rows.append(row)
    return rows


def render_spec(spec: dict, title: str) -> None:
    st.subheader(title)
    st.markdown('<div class="panel-card">', unsafe_allow_html=True)
    st.write(spec.get("summary", "No summary available."))

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Fields", len(spec.get("fields", [])))
    col2.metric("Common Rules", len(spec.get("business_rules", [])))
    col3.metric("Validations", len(spec.get("validations", [])))
    col4.metric("Calculations", len(spec.get("calculations", [])))
    st.markdown("</div>", unsafe_allow_html=True)

    (
        summary_tab,
        fields_tab,
        rules_tab,
        country_tab,
        validations_tab,
        calculations_tab,
        structure_tab,
        notes_tab,
    ) = st.tabs(
        ["Summary", "Fields", "Common Rules", "Country-Specific Rules", "Validations", "Calculations", "Structure", "Notes"]
    )

    with summary_tab:
        st.markdown("**System Summary**")
        st.write(spec.get("summary", "No summary available."))
        st.markdown("**Confidence**")
        render_table(spec.get("confidence", {}), "No confidence metadata available.")
        st.markdown("**Source Files**")
        render_table(spec.get("source_files", []), "No source files recorded.")

    with fields_tab:
        render_table(spec.get("fields", []), "No fields extracted.")

    with rules_tab:
        render_table(spec.get("business_rules", []), "No common rules extracted.")

    with country_tab:
        render_table(spec.get("country_specific_rules", []), "No country-specific rules extracted.")

    with validations_tab:
        render_table(spec.get("validations", []), "No validations extracted.")

    with calculations_tab:
        render_table(spec.get("calculations", []), "No calculations extracted.")

    with structure_tab:
        st.markdown("**UI Components**")
        render_table(spec.get("ui_components", []), "No UI components extracted.")
        ui_control_rows = build_nested_rows(
            spec.get("ui_components", []),
            "controls",
            ["component_name", "component_type", "screen_or_route"],
            child_prefix="control_",
        )
        st.markdown("**UI Controls**")
        render_table(ui_control_rows, "No UI controls extracted.")
        st.markdown("**Classes**")
        render_table(spec.get("classes", []), "No classes extracted.")
        st.markdown("**Methods**")
        render_table(spec.get("methods", []), "No methods extracted.")
        method_parameter_rows = build_nested_rows(spec.get("methods", []), "parameters", ["owner", "method_name"], child_prefix="param_")
        st.markdown("**Method Parameters**")
        render_table(method_parameter_rows, "No method parameters extracted.")
        st.markdown("**Procedures**")
        render_table(spec.get("procedures", []), "No procedures extracted.")
        procedure_parameter_rows = build_nested_rows(
            spec.get("procedures", []),
            "parameters",
            ["procedure_name", "procedure_type"],
            child_prefix="param_",
        )
        st.markdown("**Procedure Parameters**")
        render_table(procedure_parameter_rows, "No procedure parameters extracted.")
        st.markdown("**Procedure Dependencies**")
        render_table(spec.get("procedure_dependencies", []), "No procedure dependencies extracted.")
        st.markdown("**Table Dependencies**")
        render_table(spec.get("table_dependencies", []), "No table dependencies extracted.")
        st.markdown("**API Endpoints**")
        render_table(spec.get("api_endpoints", []), "No API endpoints extracted.")
        if spec.get("source_breakdown"):
            st.markdown("**Source Breakdown**")
            render_table(spec.get("source_breakdown", {}), "No source breakdown available.")

    with notes_tab:
        render_notes(spec.get("notes", []))
        errors = spec.get("read_errors", [])
        if errors:
            st.warning("File read warnings")
            render_bullets(errors, "No file warnings.")


def render_gap(gap_analysis: dict) -> None:
    confidence = to_float(gap_analysis.get("confidence", {}).get("gap_confidence", 0.0))
    missing_features = gap_analysis.get("missing_features", [])
    incorrect_implementations = gap_analysis.get("incorrect_implementations", [])
    compliance_gaps = gap_analysis.get("compliance_gaps", [])
    risks = gap_analysis.get("risks", [])
    common_rules_missed = gap_analysis.get("common_rules_missed", [])
    country_specific_rules_missed = gap_analysis.get("country_specific_rules_missed", [])

    top1, top2, top3, top4 = st.columns(4)
    with top1:
        render_metric_card("Gap Confidence", f"{confidence:.0%}", "Model-estimated certainty")
    with top2:
        render_metric_card("Missing Features", str(len(missing_features)), "Expected legacy behavior absent")
    with top3:
        render_metric_card("Compliance Gaps", str(len(compliance_gaps)), "Privacy and regulatory gaps")
    with top4:
        render_metric_card("Implementation Issues", str(len(incorrect_implementations)), "Potential parity defects")

    summary_tab, missed_rules_tab, comparison_tab, confidence_tab = st.tabs(["Gap Summary", "Rules Missed", "Rule Comparison", "Confidence"])

    with summary_tab:
        col1, col2 = st.columns(2)
        with col1:
            st.markdown('<div class="gap-card risk-medium">', unsafe_allow_html=True)
            st.markdown("**Missing Features**")
            render_bullets(missing_features, "No missing features detected.")
            st.markdown("</div>", unsafe_allow_html=True)
        with col2:
            st.markdown('<div class="gap-card risk-high">', unsafe_allow_html=True)
            st.markdown("**Compliance Gaps**")
            render_bullets(compliance_gaps, "No compliance gaps detected.")
            st.markdown("</div>", unsafe_allow_html=True)

        col3, col4 = st.columns(2)
        with col3:
            st.markdown('<div class="gap-card risk-medium">', unsafe_allow_html=True)
            st.markdown("**Incorrect Implementations**")
            render_bullets(incorrect_implementations, "No incorrect implementations detected.")
            st.markdown("</div>", unsafe_allow_html=True)
        with col4:
            st.markdown('<div class="gap-card risk-high">', unsafe_allow_html=True)
            st.markdown("**Risks**")
            if risks and all(isinstance(item, dict) for item in risks):
                render_table(risks, "No risks detected.")
            else:
                render_bullets(risks, "No risks detected.")
            st.markdown("</div>", unsafe_allow_html=True)

    with missed_rules_tab:
        common_tab, country_tab = st.tabs(["Common Rules Missed", "Country-Specific Rules Missed"])
        with common_tab:
            render_table(format_gap_rule_rows(common_rules_missed), "No common rules missed were identified.")
        with country_tab:
            render_table(format_gap_rule_rows(country_specific_rules_missed), "No country-specific rules missed were identified.")

    with comparison_tab:
        comparison_df = normalize_for_table(format_gap_rule_rows(gap_analysis.get("rule_comparison", [])))
        if not comparison_df.empty and "confidence" in comparison_df.columns:
            comparison_df["confidence"] = comparison_df["confidence"].map(
                lambda value: f"{to_float(value):.0%}" if value not in (None, "") else value
            )
        if comparison_df.empty:
            st.info("No rule comparison data available.")
        else:
            st.dataframe(sanitize_dataframe_for_display(comparison_df), use_container_width=True, hide_index=True)

    with confidence_tab:
        render_table(gap_analysis.get("confidence", {}), "No confidence data available.")


def render_logs(logs: list[dict]) -> None:
    if not logs:
        st.info("No execution logs yet.")
        return

    log_df = pd.DataFrame(logs)
    if "timestamp" in log_df.columns:
        log_df["timestamp"] = pd.to_datetime(log_df["timestamp"], errors="coerce").dt.strftime("%Y-%m-%d %H:%M:%S")

    st.markdown("**Execution Timeline**")
    for entry in logs:
        ts = entry.get("timestamp", "")
        try:
            ts = datetime.fromisoformat(ts.replace("Z", "+00:00")).strftime("%H:%M:%S")
        except ValueError:
            pass
        st.markdown(
            f"""
            <div class="panel-card" style="margin-bottom: 0.7rem;">
                <div class="small-label">{ts} | {entry.get("agent", "system")}</div>
                <div>{entry.get("message", "")}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with st.expander("Structured Log Table", expanded=False):
        st.dataframe(sanitize_dataframe_for_display(log_df), use_container_width=True, hide_index=True)


def render_flow_map(flow_map: dict, title: str) -> None:
    st.subheader(title)
    if not flow_map:
        st.info("No flow map available.")
        return

    nodes = flow_map.get("nodes", [])
    edges = flow_map.get("edges", [])
    if not nodes or not edges:
        st.info("No flow map nodes or edges available.")
        return

    dot_lines = [
        "digraph G {",
        'graph [rankdir=LR, pad="0.6", nodesep="0.7", ranksep="1.1", bgcolor="transparent"];',
        'node [shape=box, style="rounded,filled", fillcolor="#eef4fb", color="#6f94bf", fontname="Helvetica", fontsize="16", margin="0.25,0.18", penwidth="1.4"];',
        'edge [color="#4f6f96", fontname="Helvetica", fontsize="13", penwidth="1.2"];',
    ]
    for node in nodes:
        dot_lines.append(f'"{node.get("id", "")}" [label="{node.get("label", "")}"];')
    for edge in edges:
        dot_lines.append(f'"{edge.get("from", "")}" -> "{edge.get("to", "")}" [label="{edge.get("label", "")}"];')
    dot_lines.append("}")
    dot_graph = "\n".join(dot_lines)

    try:
        st.graphviz_chart(dot_graph, use_container_width=True)
    except Exception:
        st.code(dot_graph, language="dot")

    if flow_map.get("diagram_text"):
        st.caption(flow_map["diagram_text"])

    linear_flow = " -> ".join(node.get("label", "") for node in nodes if node.get("label"))
    if linear_flow:
        st.markdown("**Linear View**")
        st.code(linear_flow, language="text")

    node_tab, edge_tab = st.tabs(["Flow Nodes", "Flow Edges"])
    with node_tab:
        render_table(nodes, "No flow nodes available.")
    with edge_tab:
        render_table(edges, "No flow edges available.")


def render_step_outputs(result: dict) -> None:
    (
        legacy_code_reverse_step,
        legacy_sql_reverse_step,
        legacy_collate_step,
        target_code_reverse_step,
        target_sql_reverse_step,
        target_collate_step,
        gap_step,
    ) = st.tabs(
        [
            "Step 1: Legacy Code Reverse",
            "Step 2: Legacy SQL Reverse",
            "Step 3: Legacy Collate",
            "Step 4: Target Code Reverse",
            "Step 5: Target SQL Reverse",
            "Step 6: Target Collate",
            "Step 7: Gap Analysis",
        ]
    )

    with legacy_code_reverse_step:
        render_spec(result.get("legacy_code_reverse_spec", {}), "Legacy Source Code Reverse Engineering Output")

    with legacy_sql_reverse_step:
        render_spec(result.get("legacy_sql_reverse_spec", {}), "Legacy SQL Reverse Engineering Output")

    with legacy_collate_step:
        render_spec(result.get("legacy_spec", {}), "Legacy Consolidated Specification")
        render_flow_map(result.get("legacy_spec", {}).get("flow_map", {}), "Legacy Flow Map")

    with target_code_reverse_step:
        render_spec(result.get("target_code_reverse_spec", {}), "Target Source Code Reverse Engineering Output")

    with target_sql_reverse_step:
        render_spec(result.get("target_sql_reverse_spec", {}), "Target SQL Reverse Engineering Output")

    with target_collate_step:
        render_spec(result.get("target_spec", {}), "Target Consolidated Specification")
        render_flow_map(result.get("target_spec", {}).get("flow_map", {}), "Target Flow Map")

    with gap_step:
        render_gap(result.get("gap_analysis", {}))


def render_comparison(collated: dict, gap_analysis: dict) -> None:
    domains_tab, differences_tab, focus_tab, common_rules_tab, country_rules_tab, rules_tab = st.tabs(
        ["Shared Domains", "Key Differences", "Focus Areas", "Common Rules Missed", "Country Rules Missed", "Rule Comparison"]
    )

    with domains_tab:
        render_table(collated.get("shared_domains", []), "No shared domains available.")

    with differences_tab:
        render_table(collated.get("key_differences", []), "No key differences available.")

    with focus_tab:
        render_table(collated.get("modernization_focus_areas", []), "No focus areas available.")

    with common_rules_tab:
        render_table(format_gap_rule_rows(gap_analysis.get("common_rules_missed", [])), "No common rules missed available.")

    with country_rules_tab:
        render_table(format_gap_rule_rows(gap_analysis.get("country_specific_rules_missed", [])), "No country-specific rules missed available.")

    with rules_tab:
        render_table(format_gap_rule_rows(gap_analysis.get("rule_comparison", [])), "No rule comparison available.")


def render_flow_overview_screen() -> None:
    st.markdown(
        """
        <style>
        [data-testid="stSidebar"] {
            display: none;
        }
        .block-container {
            max-width: 100% !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    header_left, header_right = st.columns([6, 1])
    with header_right:
        if st.button("Start", type="primary", use_container_width=True):
            st.session_state["show_analysis_workspace"] = True
            st.rerun()

    steps = [
        (1, "Legacy Code Reverse"),
        (2, "Legacy SQL Reverse"),
        (3, "Legacy Collate"),
        (4, "Target Code Reverse"),
        (5, "Target SQL Reverse"),
        (6, "Target Collate"),
        (7, "Gap Analysis"),
        (8, "BRD Draft"),
        (9, "SME Approval"),
        (10, "Technical Specification"),
        (11, "Architect Approval"),
        (12, "Forward Engineering"),
        (13, "FE Proof"),
        (14, "Final Data Mapping"),
        (15, "Validation"),
    ]
    rows = [steps[i : i + 3] for i in range(0, len(steps), 3)]
    for row in rows:
        cols = st.columns(3)
        for idx, col in enumerate(cols):
            if idx >= len(row):
                col.empty()
                continue
            step_number, step_name = row[idx]
            with col:
                st.markdown(
                    f"""
                    <div class="panel-card" style="
                        min-height: 138px;
                        padding: 1.15rem 1rem;
                        display: flex;
                        flex-direction: column;
                        justify-content: space-between;
                        border-radius: 20px;
                    ">
                        <div style="
                            width: 36px;
                            height: 36px;
                            border-radius: 999px;
                            background: linear-gradient(180deg, #0b4e8c 0%, #164f7b 100%);
                            color: #ffffff;
                            display: flex;
                            align-items: center;
                            justify-content: center;
                            font-weight: 700;
                            font-size: 0.95rem;
                            margin-bottom: 0.85rem;
                        ">
                            {step_number}
                        </div>
                        <div style="
                            font-weight: 700;
                            color: #10233b;
                            font-size: 1rem;
                            line-height: 1.3;
                        ">
                            {step_name}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

    return

    steps = [
        "Legacy Code Reverse",
        "Legacy SQL Reverse",
        "Legacy Collate",
        "Target Code Reverse",
        "Target SQL Reverse",
        "Target Collate",
        "Gap Analysis",
        "BRD Draft",
        "SME Approval",
        "Technical Specification",
        "Architect Approval",
        "Forward Engineering",
        "FE Proof",
        "Final Data Mapping",
        "Validation",
    ]
    flow_cards = []
    for index, step in enumerate(steps, start=1):
        flow_cards.append(
            f"""
            <div style="
                min-width: 170px;
                max-width: 170px;
                background: rgba(255,255,255,0.96);
                border: 1px solid rgba(15,23,42,0.10);
                border-radius: 16px;
                padding: 0.9rem 0.85rem;
                box-shadow: 0 12px 28px rgba(15, 23, 42, 0.08);
                text-align: center;
                flex: 0 0 auto;
            ">
                <div style="
                    width: 32px;
                    height: 32px;
                    border-radius: 999px;
                    background: #0b4e8c;
                    color: white;
                    display: inline-flex;
                    align-items: center;
                    justify-content: center;
                    font-weight: 700;
                    margin-bottom: 0.55rem;
                ">{index}</div>
                <div style="
                    color: #10233b;
                    font-weight: 700;
                    font-size: 0.92rem;
                    line-height: 1.25;
                ">{step}</div>
            </div>
            """
        )
        if index < len(steps):
            flow_cards.append(
                """
                <div style="
                    flex: 0 0 auto;
                    align-self: center;
                    color: #5e7f9a;
                    font-size: 1.5rem;
                    font-weight: 700;
                    padding: 0 0.35rem;
                ">→</div>
                """
            )

    st.markdown(
        f"""
        <div style="
            display: flex;
            gap: 0.2rem;
            overflow-x: auto;
            padding: 0.4rem 0.1rem 1rem 0.1rem;
            align-items: stretch;
        ">
            {''.join(flow_cards)}
        </div>
        """,
        unsafe_allow_html=True,
    )

    if st.button("Go to Analysis", type="primary", use_container_width=True):
        st.session_state["show_analysis_workspace"] = True
        st.rerun()


def render_flow_overview_screen_v2() -> None:
    st.markdown(
        """
        <style>
        [data-testid="stSidebar"] {
            display: none;
        }
        .block-container {
            max-width: 100% !important;
        }
        .overview-hero {
            border-radius: 28px;
            padding: 1.65rem 1.75rem;
            background:
                radial-gradient(circle at top right, rgba(11, 78, 140, 0.24), transparent 30%),
                radial-gradient(circle at bottom left, rgba(47, 128, 88, 0.18), transparent 24%),
                linear-gradient(135deg, #f9fbff 0%, #edf5ff 52%, #f5fbf8 100%);
            border: 1px solid rgba(15, 23, 42, 0.08);
            box-shadow: 0 24px 42px rgba(15, 23, 42, 0.08);
            margin-bottom: 1rem;
        }
        .overview-title {
            font-size: 2.3rem;
            line-height: 1.06;
            color: #10233b;
            font-weight: 800;
            margin: 0.35rem 0 0.55rem 0;
        }
        .overview-copy {
            color: #53677d;
            font-size: 1rem;
            line-height: 1.48;
            max-width: 880px;
        }
        .overview-chip {
            display: inline-flex;
            align-items: center;
            padding: 0.38rem 0.74rem;
            border-radius: 999px;
            background: rgba(11, 78, 140, 0.09);
            color: #0b4e8c;
            font-size: 0.82rem;
            font-weight: 700;
            margin-right: 0.45rem;
            margin-top: 0.6rem;
        }
        .overview-band {
            border-radius: 22px;
            padding: 1rem 1.08rem;
            background: rgba(255, 255, 255, 0.86);
            border: 1px solid rgba(15, 23, 42, 0.08);
            box-shadow: 0 16px 30px rgba(15, 23, 42, 0.05);
            margin-bottom: 1rem;
        }
        .overview-step-card {
            min-height: 228px;
            border-radius: 24px;
            padding: 1.15rem 1.05rem 1rem 1.05rem;
            color: #10233b;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            box-shadow: 0 18px 32px rgba(15, 23, 42, 0.07);
            border: 1px solid rgba(15, 23, 42, 0.08);
        }
        .overview-step-index {
            width: 38px;
            height: 38px;
            border-radius: 999px;
            background: rgba(16, 35, 59, 0.93);
            color: #ffffff;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            font-weight: 800;
            font-size: 0.95rem;
        }
        .overview-step-name {
            font-size: 1.02rem;
            font-weight: 800;
            line-height: 1.24;
            margin-top: 0.85rem;
        }
        .overview-step-purpose {
            color: #45596f;
            font-size: 0.91rem;
            line-height: 1.42;
            margin-top: 0.55rem;
            min-height: 64px;
        }
        .overview-step-artifact {
            margin-top: 0.8rem;
            border-radius: 16px;
            background: rgba(255, 255, 255, 0.76);
            padding: 0.7rem 0.78rem;
        }
        .overview-step-label {
            font-size: 0.72rem;
            text-transform: uppercase;
            letter-spacing: 0.07em;
            color: #5f738a;
            font-weight: 800;
            margin-bottom: 0.18rem;
        }
        .overview-step-value {
            color: #10233b;
            font-size: 0.9rem;
            font-weight: 700;
            line-height: 1.32;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    hero_left, hero_right = st.columns([5, 1.2])
    with hero_left:
        st.markdown(
            """
            <div class="overview-hero">
                <div class="eyebrow">AI Migration Journey</div>
                <div class="overview-title">Understand the full modernization flow in one glance</div>
                <div class="overview-copy">
                    The platform downloads legacy and target code from GitHub, combines that with local SQL analysis, builds approved migration artifacts, generates forward-engineered code and automated tests, validates the result, and finally opens a draft PR with AI-assisted review.
                </div>
                <div>
                    <span class="overview-chip">GitHub Code Snapshots</span>
                    <span class="overview-chip">Local SQL Reverse Engineering</span>
                    <span class="overview-chip">Approval Gates</span>
                    <span class="overview-chip">Generated Code + Tests</span>
                    <span class="overview-chip">Draft PR + AI Review</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with hero_right:
        if st.button("Start Workspace", type="primary", use_container_width=True):
            st.session_state["show_analysis_workspace"] = True
            st.rerun()

    st.markdown(
        """
        <div class="overview-band">
            <div class="section-title">Sequential Stages And Primary Artifacts</div>
            <div class="muted-copy">Each step below shows its purpose and its main output artifact, so someone new to the platform can immediately understand the sequence and what gets produced at each milestone.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    steps = [
        {"number": 1, "name": "Legacy Code Snapshot", "purpose": "Download legacy application code from GitHub main into a local working snapshot before reverse engineering.", "artifact": "Local legacy code snapshot", "tone": "linear-gradient(135deg, #f4f9ff 0%, #e5f0ff 100%)"},
        {"number": 2, "name": "Legacy SQL Reverse", "purpose": "Analyze the local legacy SQL folder to capture procedures, data structures, and rule behavior.", "artifact": "Legacy SQL reverse spec", "tone": "linear-gradient(135deg, #f6fbff 0%, #e8f6ff 100%)"},
        {"number": 3, "name": "Legacy Canonical Spec", "purpose": "Collate legacy code and SQL findings into one business-readable baseline specification.", "artifact": "Legacy consolidated specification", "tone": "linear-gradient(135deg, #f6fff8 0%, #e7f8ed 100%)"},
        {"number": 4, "name": "Target Code Snapshot", "purpose": "Download target application code from GitHub main into a local snapshot used for reverse engineering and generation.", "artifact": "Local target code snapshot", "tone": "linear-gradient(135deg, #fff8f1 0%, #ffefdb 100%)"},
        {"number": 5, "name": "Target SQL Reverse", "purpose": "Analyze the local target SQL folder separately so structural and rule comparisons stay accurate.", "artifact": "Target SQL reverse spec", "tone": "linear-gradient(135deg, #fff9f4 0%, #fff1e4 100%)"},
        {"number": 6, "name": "Target Canonical Spec", "purpose": "Build the target-side specification and flow view from downloaded code plus local SQL.", "artifact": "Target consolidated specification", "tone": "linear-gradient(135deg, #fffaf6 0%, #fbeedf 100%)"},
        {"number": 7, "name": "Gap Analysis", "purpose": "Identify missing features, compliance issues, parity gaps, and migration risks across the two systems.", "artifact": "Gap and risk assessment", "tone": "linear-gradient(135deg, #fff6f6 0%, #ffe8e8 100%)"},
        {"number": 8, "name": "BRD Draft", "purpose": "Turn the gap findings into a migration-ready business requirements draft.", "artifact": "Requirements / BRD draft", "tone": "linear-gradient(135deg, #f8f6ff 0%, #eee8ff 100%)"},
        {"number": 9, "name": "SME Approval", "purpose": "Pause for business review so the requirements become the approved migration baseline.", "artifact": "Approved requirements", "tone": "linear-gradient(135deg, #f3fbff 0%, #def2ff 100%)"},
        {"number": 10, "name": "Technical Specification", "purpose": "Generate the target design directly from the approved migration requirements.", "artifact": "Technical specification draft", "tone": "linear-gradient(135deg, #f6fffb 0%, #e3f8ef 100%)"},
        {"number": 11, "name": "Architect Approval", "purpose": "Pause for architectural review before generation continues.", "artifact": "Approved technical design", "tone": "linear-gradient(135deg, #f8fcff 0%, #e8f2ff 100%)"},
        {"number": 12, "name": "Forward Engineering", "purpose": "Generate target code and automated unit tests on top of the downloaded target folder structure.", "artifact": "Generated code + unit test assets", "tone": "linear-gradient(135deg, #f4fff9 0%, #dbf6e8 100%)"},
        {"number": 13, "name": "Forward Engineering Proof", "purpose": "Show added and modified files so reviewers can inspect the generated target candidate before publishing.", "artifact": "Generated diff and proof view", "tone": "linear-gradient(135deg, #fffdf5 0%, #fff5dc 100%)"},
        {"number": 14, "name": "Final Data Mapping", "purpose": "Measure mapping coverage, reconciliation rules, unmapped fields, and migration data risks.", "artifact": "Data mapping and reconciliation report", "tone": "linear-gradient(135deg, #f7fbff 0%, #e6efff 100%)"},
        {"number": 15, "name": "Validation", "purpose": "Validate generated target artifacts against source behavior and approved design using the local workspace.", "artifact": "Validation findings and sync score", "tone": "linear-gradient(135deg, #fff8fb 0%, #ffeaf5 100%)"},
        {"number": 16, "name": "AI Code Review", "purpose": "Run an AI-assisted review over the generated diff and prepare detailed review findings for the PR.", "artifact": "AI review findings", "tone": "linear-gradient(135deg, #f5f8ff 0%, #e7ebff 100%)"},
        {"number": 17, "name": "Draft PR", "purpose": "Create a new target branch from main, publish only code and test files, and open a draft PR back to main.", "artifact": "Draft GitHub PR", "tone": "linear-gradient(135deg, #f4fffd 0%, #dff6ef 100%)"},
    ]
    rows = [steps[i : i + 4] for i in range(0, len(steps), 4)]
    for row in rows:
        cols = st.columns(4)
        for idx, col in enumerate(cols):
            if idx >= len(row):
                col.empty()
                continue
            step = row[idx]
            with col:
                st.markdown(
                    f"""
                    <div class="overview-step-card" style="background: {step['tone']};">
                        <div>
                            <div class="overview-step-index">{step['number']}</div>
                            <div class="overview-step-name">{step['name']}</div>
                            <div class="overview-step-purpose">{step['purpose']}</div>
                        </div>
                        <div class="overview-step-artifact">
                            <div class="overview-step-label">Primary Artifact</div>
                            <div class="overview-step-value">{step['artifact']}</div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )


def main() -> None:
    apply_enterprise_theme()
    init_approval_state()
    st.title("AI Powered Migration Platform")
    #st.caption("Reverse engineer legacy and target insurance systems, collate structured specs, and highlight migration gaps.")

    default_legacy_repo = "https://github.com/arkadeep87/SourceVBApp.git"
    default_legacy_branch = "main"
    default_legacy_code_root = ""
    default_legacy_code = str(settings.outputs_dir / "github_snapshots" / "legacy_code")
    default_legacy_sql = str(settings.sample_inputs_dir / "legacy" / "quote_generation" / "sql")
    default_target_repo = "https://github.com/arkadeep87/Target-App.git"
    default_target_branch = "main"
    default_target_code_root = ""
    default_target_code = str(settings.outputs_dir / "github_snapshots" / "target_code")
    default_target_sql = str(settings.sample_inputs_dir / "target" / "quote_generation" / "sql")
    default_legacy_data_dictionary = str(settings.sample_inputs_dir / "data_dictionary" / "legacy")
    default_target_data_dictionary = str(settings.sample_inputs_dir / "data_dictionary" / "target")
    model_options = MODEL_OPTIONS.copy()
    if settings.model and settings.model not in model_options:
        model_options.insert(0, settings.model)
    selected_model_index = model_options.index(settings.model) if settings.model in model_options else 0

    with st.sidebar:
        st.markdown("## Control Center")
        st.caption("Configure inputs, start the workflow, and review runtime settings.")
        st.markdown("### Legacy Inputs")
        legacy_repo_url = st.text_input(
            "Legacy GitHub repo",
            value=default_legacy_repo,
            help="GitHub repository URL used to fetch legacy/source code before analysis. The selected branch is downloaded into the local snapshot folder below.",
        )
        legacy_repo_branch = st.text_input(
            "Legacy branch",
            value=default_legacy_branch,
            help="Branch to download from the legacy/source repository.",
        )
        legacy_repo_code_root = default_legacy_code_root
        legacy_code_folder = default_legacy_code
        legacy_sql_folder = st.text_input(
            "Legacy SQL folder",
            value=default_legacy_sql,
            help="Local legacy SQL folder. SQL stays local and is not fetched from GitHub or included in PR publishing.",
        )
        st.markdown("### Target Inputs")
        target_repo_url = st.text_input(
            "Target GitHub repo",
            value=default_target_repo,
            help="GitHub repository URL used to fetch target code and later publish the generated PR.",
        )
        target_repo_branch = st.text_input(
            "Target branch",
            value=default_target_branch,
            help="Branch to download from and later use as the PR base branch.",
        )
        target_repo_code_root = default_target_code_root
        target_code_folder = default_target_code
        target_sql_folder = st.text_input(
            "Target SQL folder",
            value=default_target_sql,
            help="Local target SQL folder. SQL stays local and is excluded from GitHub PR publishing.",
        )
        st.markdown("### Data Dictionaries")
        legacy_data_dictionary_folder = st.text_input(
            "Legacy data dictionary folder",
            value=default_legacy_data_dictionary,
            help="Provide a folder containing the initial legacy data dictionary files such as CSV, JSON, TXT, or Markdown.",
        )
        target_data_dictionary_folder = st.text_input(
            "Initial target data dictionary folder",
            value=default_target_data_dictionary,
            help="Provide a folder containing the initial target data dictionary files used before forward engineering.",
        )
        selected_model = st.selectbox(
            "Model",
            options=model_options,
            index=selected_model_index,
            format_func=lambda model_id: MODEL_LABELS.get(model_id, model_id),
        )
        cache_enabled = st.toggle(
            "Caching",
            value=settings.cache_enabled,
            help="Controls cache reuse for reverse engineering, gap analysis, requirements, technical specification, forward engineering, and generated target-candidate outputs.",
        )
        run_clicked = st.button("Run Modernization Analysis", type="primary", use_container_width=True)
        st.markdown("### Workflow")
        st.caption("GitHub code is downloaded into local snapshot folders first. SQL remains local. PR publishing later includes code and generated test files, but excludes SQL.")
        st.markdown("- Download legacy code from GitHub into a local snapshot")
        st.markdown("- Download target code from GitHub into a local snapshot")
        st.markdown("- Reverse legacy code and local legacy SQL")
        st.markdown("- Collate the legacy canonical specification")
        st.markdown("- Reverse target code and local target SQL")
        st.markdown("- Collate the target canonical specification")
        st.markdown("- Build flow maps and run gap analysis")
        st.markdown("- Generate the BRD draft and wait for SME approval")
        st.markdown("- Generate the technical specification and wait for architect approval")
        st.markdown("- Generate forward-engineered code and automated unit tests")
        st.markdown("- Review forward-engineering proof, validation, and final data mapping")
        st.markdown("- Run AI Code Review and post detailed review comments to the PR")
        st.markdown("- Raise a draft GitHub PR with code and test files only")
        if st.session_state.get("analysis_result"):
            st.success("Latest analysis is loaded in the workspace.")

    legacy_repo_full_name = parse_github_repo_full_name(legacy_repo_url)
    target_repo_full_name = parse_github_repo_full_name(target_repo_url)
    settings.model = selected_model
    settings.cache_enabled = cache_enabled
    if not st.session_state.get("show_analysis_workspace"):
        render_flow_overview_screen_v2()
        return
    auto_run_analysis = st.session_state.pop("auto_run_analysis", False)
    auto_run_stage = st.session_state.pop("auto_run_stage", "")

    intro_col, stat_col = st.columns([2, 1])
    with intro_col:
        st.markdown(
            """
            <div class="hero-card">
                <div class="eyebrow">Intelligent Migration</div>
                <div class="hero-title">AI driven - Human assisted platform for application migration to modern platforms</div>
                <div class="hero-copy">
                    Reverse Engineering | Canonicalization | Gap Analysis | Business Requirement Creation | Tech Spec Generation | Forward Engineering
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with stat_col:
        st.markdown(
            """
            <div class="panel-card">
                <div class="section-title">Launch Ready</div>
                <div class="muted-copy">Code is fetched from GitHub into local snapshots for analysis and generation. SQL remains local for reverse engineering, validation, and mapping.</div>
                <div style="margin-top: 0.85rem;">
                    <span class="status-chip">Legacy VB.NET</span>
                    <span class="status-chip">Sybase SQL</span>
                    <span class="status-chip">Angular</span>
                    <span class="status-chip">Node.js</span>
                    <span class="status-chip">PostgreSQL</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    if run_clicked or auto_run_analysis:
        st.session_state["active_main_section"] = "Migration Dashboard"
        approved_requirements_for_run = st.session_state.get("approved_requirements")
        approved_technical_spec_for_run = st.session_state.get("approved_technical_spec")
        prior_state_for_run = st.session_state.get("analysis_result") if auto_run_analysis else None
        legacy_repo_full_name = parse_github_repo_full_name(legacy_repo_url)
        target_repo_full_name = parse_github_repo_full_name(target_repo_url)
        runtime_legacy_code_folder = legacy_code_folder
        runtime_target_code_folder = target_code_folder

        if run_clicked and not auto_run_analysis:
            approved_requirements_for_run = None
            approved_technical_spec_for_run = None
            st.session_state["approved_requirements"] = None
            st.session_state["approved_technical_spec"] = None
            st.session_state["sme_comments"] = ""
            st.session_state["architect_comments"] = ""
            st.session_state["validation_result"] = None
            st.session_state["validation_error"] = None
            st.session_state["data_mapping_result"] = None
            st.session_state["data_mapping_error"] = None
            st.session_state["ai_code_review_result"] = None
            st.session_state["ai_code_review_error"] = None
            st.session_state["github_publish_result"] = None
            st.session_state["github_publish_error"] = None
            st.session_state["github_branch_name"] = ""
            st.session_state["github_commit_message"] = ""
            st.session_state["github_pr_title"] = ""

        st.session_state.pop("analysis_error", None)
        st.session_state.pop("analysis_trace_dir", None)
        progress_placeholder = st.empty()
        status_placeholder = st.empty()
        live_dashboard_placeholder = st.empty()
        live_logs_placeholder = st.empty()
        trace_dir = create_run_trace_dir()
        st.session_state["analysis_trace_dir"] = str(trace_dir)
        try:
            if not legacy_repo_full_name:
                raise RuntimeError("Unable to parse the legacy GitHub repository URL.")
            if not target_repo_full_name:
                raise RuntimeError("Unable to parse the target GitHub repository URL.")
            status_placeholder.info("Downloading legacy and target code from GitHub into local snapshots.")
            runtime_legacy_code_folder = github_fetch_repo_code_snapshot(
                repo_full_name=legacy_repo_full_name,
                branch=legacy_repo_branch,
                code_root=legacy_repo_code_root,
                destination=legacy_code_folder,
            )
            runtime_target_code_folder = github_fetch_repo_code_snapshot(
                repo_full_name=target_repo_full_name,
                branch=target_repo_branch,
                code_root=target_repo_code_root,
                destination=target_code_folder,
            )
            progress_bar = progress_placeholder.progress(0, text="Initializing modernization workflow...")
            status_placeholder.info("Preparing artifact discovery and analysis context.")
            write_json_trace(
                trace_dir / "run_context.json",
                {
                    "started_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC"),
                    "legacy_repo": legacy_repo_full_name,
                    "legacy_branch": legacy_repo_branch,
                    "legacy_code_folder": runtime_legacy_code_folder,
                    "legacy_sql_folder": legacy_sql_folder,
                    "target_repo": target_repo_full_name,
                    "target_branch": target_repo_branch,
                    "target_code_folder": runtime_target_code_folder,
                    "target_sql_folder": target_sql_folder,
                    "model": settings.model,
                    "cache_enabled": settings.cache_enabled,
                },
            )
            result: dict = {}
            if auto_run_analysis and auto_run_stage == "technical_spec" and approved_requirements_for_run:
                workflow_stream = stream_resume_from_requirements_approval(
                    prior_state=prior_state_for_run,
                    approved_requirements=approved_requirements_for_run,
                )
            elif (
                auto_run_analysis
                and auto_run_stage == "forward_engineering"
                and approved_requirements_for_run
                and approved_technical_spec_for_run
            ):
                workflow_stream = stream_resume_from_technical_spec_approval(
                    prior_state=prior_state_for_run,
                    approved_requirements=approved_requirements_for_run,
                    approved_technical_spec=approved_technical_spec_for_run,
                )
            else:
                workflow_stream = stream_workflow(
                    legacy_code_folder=runtime_legacy_code_folder,
                    legacy_sql_folder=legacy_sql_folder,
                    target_code_folder=runtime_target_code_folder,
                    target_sql_folder=target_sql_folder,
                    approved_requirements=approved_requirements_for_run,
                    approved_technical_spec=approved_technical_spec_for_run,
                    prior_state=prior_state_for_run,
                )

            for partial_state in workflow_stream:
                result = partial_state
                write_run_trace_snapshot(trace_dir, partial_state, "running")
                logs = partial_state.get("logs", [])
                with live_dashboard_placeholder.container():
                    st.markdown("### Live Migration Dashboard")
                    render_migration_dashboard(partial_state, live=True)
                render_live_execution_status(live_logs_placeholder, logs)
                seen = {entry.get("agent") for entry in logs}
                completed = sum(1 for key, _label in WORKFLOW_PHASES if key in seen)
                progress_value = completed / len(WORKFLOW_PHASES)
                current_label = "Building execution plan"
                for key, label in WORKFLOW_PHASES:
                    if key not in seen:
                        current_label = label
                        break
                if logs:
                    status_placeholder.info(logs[-1].get("message", current_label))
                progress_bar.progress(progress_value, text=f"{current_label}...")

            st.session_state["analysis_result"] = result
            sync_approval_state_from_result(result)
            finalize_run_trace(trace_dir, result)
            live_dashboard_placeholder.empty()
            live_logs_placeholder.empty()
            progress_bar.progress(100, text="Analysis complete.")
            status_placeholder.success(f"Modernization analysis completed successfully. Trace saved to {trace_dir}")
        except Exception as exc:
            st.session_state["analysis_error"] = str(exc)
            write_error_trace(trace_dir, str(exc), result if "result" in locals() else {})
            live_dashboard_placeholder.empty()
            live_logs_placeholder.empty()
            status_placeholder.error(f"Analysis failed. Trace saved to {trace_dir}")

    if st.session_state.get("analysis_error"):
        st.error(st.session_state["analysis_error"])
    if st.session_state.get("analysis_trace_dir"):
        st.caption(f"Run trace folder: {st.session_state['analysis_trace_dir']}")

    result = st.session_state.get("analysis_result")
    if not result:
        st.markdown("### Ready to analyze")
        st.write("Use the sample folders or point the app at separate source-code and SQL folders for legacy and target systems.")
        return

    sync_approval_state_from_result(result)

    requirements_draft = deep_copy_document(result.get("requirements_draft", {}))
    technical_spec_draft = deep_copy_document(result.get("technical_spec_draft", {}))
    approved_requirements = st.session_state.get("approved_requirements")
    approved_technical_spec = st.session_state.get("approved_technical_spec")

    if approved_requirements:
        result["approved_requirements"] = approved_requirements
        result["requirements_draft"] = approved_requirements
        requirements_draft = deep_copy_document(approved_requirements)
    if approved_technical_spec:
        result["approved_technical_spec"] = approved_technical_spec
        result["technical_spec_draft"] = approved_technical_spec
        technical_spec_draft = deep_copy_document(approved_technical_spec)

    persisted_validation_result = result.get("validation_result")
    if st.session_state.get("validation_result") is None and persisted_validation_result:
        st.session_state["validation_result"] = persisted_validation_result
    persisted_data_mapping_result = result.get("data_mapping_result")
    if st.session_state.get("data_mapping_result") is None and persisted_data_mapping_result:
        st.session_state["data_mapping_result"] = persisted_data_mapping_result
    persisted_ai_code_review_result = result.get("ai_code_review_result")
    if st.session_state.get("ai_code_review_result") is None and persisted_ai_code_review_result:
        st.session_state["ai_code_review_result"] = persisted_ai_code_review_result

    workflow_excel_bytes = build_workflow_excel_export(result)

    requirements_status = get_approval_status(requirements_draft, PENDING_SME_APPROVAL)
    technical_spec_status = get_approval_status(technical_spec_draft, PENDING_ARCHITECT_APPROVAL)
    generated_target = {"generated_root": "", "generated_files": []}
    forward_comparison_items: list[dict] = []
    if result.get("forward_engineering_output"):
        generated_target = materialize_forward_engineering_output(
            result.get("forward_engineering_output", {}),
            target_code_folder,
            target_sql_folder,
        )
        forward_comparison_items = build_forward_engineering_comparison(
            generated_target.get("generated_files", []),
            target_code_folder,
            target_sql_folder,
        )
        publish_defaults = build_publish_defaults(result, target_repo_branch)
        if not st.session_state.get("github_branch_name"):
            st.session_state["github_branch_name"] = publish_defaults["branch_name"]
        if not st.session_state.get("github_base_branch"):
            st.session_state["github_base_branch"] = publish_defaults["base_branch"]
        if not st.session_state.get("github_commit_message"):
            st.session_state["github_commit_message"] = publish_defaults["commit_message"]
        if not st.session_state.get("github_pr_title"):
            st.session_state["github_pr_title"] = publish_defaults["pr_title"]

    section_options = [
        "Migration Dashboard",
        "Overview",
        "Step Outputs",
        "Legacy Spec",
        "Target Spec",
        "Flow Maps",
        "Comparison",
        "Gap Analysis",
        "Requirements Draft",
        "Technical Specification Draft",
        "Forward Engineering",
        "Forward Engineering Proof",
        "Final Data Mapping",
        "Validation",
        "AI Code Review",
        "Execution Logs",
        "Raw JSON",
    ]
    pending_main_section = st.session_state.pop("pending_main_section", "")
    if pending_main_section in section_options:
        st.session_state["active_main_section"] = pending_main_section
    st.session_state.setdefault("active_main_section", "Migration Dashboard")
    st.markdown("**Workspace Sections**")
    if openpyxl is None:
        st.info("Install `openpyxl` to enable the Excel export of workflow outputs.")
    elif workflow_excel_bytes:
        st.download_button(
            label="Download Workflow Excel",
            data=workflow_excel_bytes,
            file_name="modernization_workflow_export.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_workflow_excel",
            use_container_width=False,
        )
    active_section = st.session_state.get("active_main_section", "Migration Dashboard")
    if active_section not in section_options:
        active_section = "Migration Dashboard"
        st.session_state["active_main_section"] = active_section

    if active_section != "Migration Dashboard":
        nav_left, nav_right = st.columns([1, 5])
        with nav_left:
            if st.button("Back to Dashboard", use_container_width=True):
                st.session_state["active_main_section"] = "Migration Dashboard"
                st.rerun()
        with nav_right:
            st.caption(f"Detail View: {active_section}")

    if active_section == "Migration Dashboard":
        dashboard_state = dict(result)
        dashboard_state["validation_result"] = st.session_state.get("validation_result") or result.get("validation_result")
        dashboard_state["data_mapping_result"] = st.session_state.get("data_mapping_result") or result.get("data_mapping_result")
        dashboard_state["ai_code_review_result"] = st.session_state.get("ai_code_review_result") or result.get("ai_code_review_result")
        dashboard_state["generated_target"] = generated_target
        dashboard_state["forward_comparison_items"] = forward_comparison_items
        render_migration_dashboard(dashboard_state)

    elif active_section == "Overview":
        gap_analysis = result.get("gap_analysis", {})
        collated = result.get("collated_spec", {})
        overview_gap_confidence = to_float(gap_analysis.get("confidence", {}).get("gap_confidence", 0.0))
        top1, top2, top3 = st.columns(3)
        with top1:
            render_metric_card("Missing Features", str(len(gap_analysis.get("missing_features", []))), "Feature parity shortfalls")
        with top2:
            render_metric_card("Compliance Gaps", str(len(gap_analysis.get("compliance_gaps", []))), "Regulatory exposures detected")
        with top3:
            render_metric_card("Gap Confidence", f"{overview_gap_confidence:.0%}", "Confidence in assessment quality")

        st.markdown("**Modernization Focus Areas**")
        render_table(collated.get("modernization_focus_areas", []), "No focus areas available.")
        st.markdown("**Key Differences**")
        render_table(collated.get("key_differences", []), "No key differences available.")

    elif active_section == "Step Outputs":
        render_step_outputs(result)

    elif active_section == "Legacy Spec":
        render_spec(result.get("legacy_spec", {}), "Legacy Insurance System")

    elif active_section == "Target Spec":
        render_spec(result.get("target_spec", {}), "Target Insurance System")

    elif active_section == "Flow Maps":
        legacy_flow_tab, target_flow_tab = st.tabs(["Legacy Flow Map", "Target Flow Map"])
        with legacy_flow_tab:
            render_flow_map(result.get("legacy_spec", {}).get("flow_map", {}), "Legacy System Flow Map")
        with target_flow_tab:
            render_flow_map(result.get("target_spec", {}).get("flow_map", {}), "Target System Flow Map")

    elif active_section == "Comparison":
        render_comparison(result.get("collated_spec", {}), result.get("gap_analysis", {}))

    elif active_section == "Gap Analysis":
        render_gap(result.get("gap_analysis", {}))

    elif active_section == "Requirements Draft":
        render_requirements_document(requirements_draft)
        st.markdown("**SME Review Controls**")
        render_approval_summary("Requirements", requirements_status, st.session_state.get("sme_comments", ""))
        st.text_area(
            "SME comments",
            key="sme_comments",
            height=120,
            placeholder="Capture SME review notes for the requirements draft.",
        )
        approve_col, reject_col = st.columns(2)
        with approve_col:
            if st.button("Approve Requirements", use_container_width=True):
                if requirements_draft:
                    approved_doc = update_document_approval(
                        requirements_draft,
                        SME_APPROVED,
                        st.session_state.get("sme_comments", ""),
                        "SME",
                    )
                    st.session_state["approved_requirements"] = approved_doc
                    st.session_state["approved_technical_spec"] = None
                    st.session_state["analysis_result"]["requirements_draft"] = approved_doc
                    st.session_state["auto_run_analysis"] = True
                    st.session_state["auto_run_stage"] = "technical_spec"
                    st.session_state["validation_result"] = None
                    st.session_state["validation_error"] = None
                    st.session_state["data_mapping_result"] = None
                    st.session_state["data_mapping_error"] = None
                    st.session_state["ai_code_review_result"] = None
                    st.session_state["ai_code_review_error"] = None
                    st.session_state["github_publish_result"] = None
                    st.session_state["github_publish_error"] = None
                    st.session_state["github_branch_name"] = ""
                    st.session_state["github_commit_message"] = ""
                    st.session_state["github_pr_title"] = ""
                    st.success("Requirements approved. Continuing automatically to technical specification generation.")
                    st.rerun()
                else:
                    st.warning("No requirements draft is available yet.")
        with reject_col:
            if st.button("Reject Requirements", use_container_width=True):
                if requirements_draft:
                    rejected_doc = update_document_approval(
                        requirements_draft,
                        SME_REJECTED,
                        st.session_state.get("sme_comments", ""),
                        "SME",
                    )
                    st.session_state["approved_requirements"] = None
                    st.session_state["approved_technical_spec"] = None
                    st.session_state["analysis_result"]["requirements_draft"] = rejected_doc
                    st.session_state["ai_code_review_result"] = None
                    st.session_state["ai_code_review_error"] = None
                    st.session_state["github_publish_result"] = None
                    st.session_state["github_publish_error"] = None
                    st.session_state["github_branch_name"] = ""
                    st.session_state["github_commit_message"] = ""
                    st.session_state["github_pr_title"] = ""
                    st.warning("Requirements rejected. Run the analysis again to regenerate the draft.")
                    st.rerun()
                else:
                    st.warning("No requirements draft is available yet.")

    elif active_section == "Technical Specification Draft":
        if requirements_status != SME_APPROVED:
            st.info("Technical specification is blocked until the requirements draft is SME approved.")
        else:
            render_technical_spec_document(technical_spec_draft)

        st.markdown("**Architect Review Controls**")
        render_approval_summary("Technical Specification", technical_spec_status, st.session_state.get("architect_comments", ""))
        st.text_area(
            "Architect comments",
            key="architect_comments",
            height=120,
            placeholder="Capture architect review notes for the technical specification draft.",
        )
        approve_arch_col, reject_arch_col = st.columns(2)
        with approve_arch_col:
            if st.button("Approve Technical Spec", use_container_width=True):
                if requirements_status != SME_APPROVED:
                    st.warning("Approve the requirements draft first.")
                elif technical_spec_draft:
                    approved_doc = update_document_approval(
                        technical_spec_draft,
                        ARCHITECT_APPROVED,
                        st.session_state.get("architect_comments", ""),
                        "Architect",
                    )
                    st.session_state["approved_technical_spec"] = approved_doc
                    st.session_state["analysis_result"]["technical_spec_draft"] = approved_doc
                    st.session_state["auto_run_analysis"] = True
                    st.session_state["auto_run_stage"] = "forward_engineering"
                    st.session_state["validation_result"] = None
                    st.session_state["validation_error"] = None
                    st.session_state["data_mapping_result"] = None
                    st.session_state["data_mapping_error"] = None
                    st.session_state["ai_code_review_result"] = None
                    st.session_state["ai_code_review_error"] = None
                    st.session_state["github_publish_result"] = None
                    st.session_state["github_publish_error"] = None
                    st.session_state["github_branch_name"] = ""
                    st.session_state["github_commit_message"] = ""
                    st.session_state["github_pr_title"] = ""
                    st.success("Technical specification approved. Continuing automatically to forward engineering generation.")
                    st.rerun()
                else:
                    st.warning("No technical specification draft is available yet. Run the analysis after SME approval.")
        with reject_arch_col:
            if st.button("Reject Technical Spec", use_container_width=True):
                if technical_spec_draft:
                    rejected_doc = update_document_approval(
                        technical_spec_draft,
                        ARCHITECT_REJECTED,
                        st.session_state.get("architect_comments", ""),
                        "Architect",
                    )
                    st.session_state["approved_technical_spec"] = None
                    st.session_state["analysis_result"]["technical_spec_draft"] = rejected_doc
                    st.session_state["ai_code_review_result"] = None
                    st.session_state["ai_code_review_error"] = None
                    st.session_state["github_publish_result"] = None
                    st.session_state["github_publish_error"] = None
                    st.session_state["github_branch_name"] = ""
                    st.session_state["github_commit_message"] = ""
                    st.session_state["github_pr_title"] = ""
                    st.warning("Technical specification rejected. Run the analysis again to regenerate the draft.")
                    st.rerun()
                else:
                    st.warning("No technical specification draft is available yet.")

    elif active_section == "Forward Engineering":
        if requirements_status != SME_APPROVED:
            st.info("Forward engineering is blocked until the requirements draft is SME approved.")
        elif technical_spec_status != ARCHITECT_APPROVED:
            st.info("Forward engineering is blocked until the technical specification is architect approved.")
        else:
            render_forward_engineering_summary(
                result.get("forward_engineering_output", {}),
                forward_comparison_items,
                generated_target.get("generated_root", ""),
            )

    elif active_section == "Forward Engineering Proof":
        if requirements_status != SME_APPROVED:
            st.info("Forward engineering proof is blocked until the requirements draft is SME approved.")
        elif technical_spec_status != ARCHITECT_APPROVED:
            st.info("Forward engineering proof is blocked until the technical specification is architect approved.")
        else:
            render_forward_engineering_proof(forward_comparison_items)
            st.markdown("**Post-Generation Analysis**")
            if st.button("Validate Generated Target", use_container_width=True):
                try:
                    run_validation_from_workspace(
                        result=result,
                        requirements_draft=requirements_draft,
                        technical_spec_draft=technical_spec_draft,
                        generated_target=generated_target,
                    )
                except Exception as exc:
                    st.session_state["validation_error"] = str(exc)
                    st.rerun()
            if st.button("Data Mapping and Reconciliation", use_container_width=True):
                try:
                    run_final_data_mapping_from_workspace(
                        result=result,
                        generated_target=generated_target,
                        legacy_data_dictionary_folder=legacy_data_dictionary_folder,
                        target_data_dictionary_folder=target_data_dictionary_folder,
                    )
                except Exception as exc:
                    st.session_state["data_mapping_error"] = str(exc)
                    st.rerun()

    elif active_section == "Final Data Mapping":
        if requirements_status != SME_APPROVED:
            st.info("Final data mapping is blocked until the requirements draft is SME approved.")
        elif technical_spec_status != ARCHITECT_APPROVED:
            st.info("Final data mapping is blocked until the technical specification is architect approved.")
        elif not result.get("forward_engineering_output"):
            st.info("Final data mapping is available after forward engineering generates target artifacts.")
        else:
            if not (st.session_state.get("data_mapping_result") or result.get("data_mapping_result")):
                st.markdown("**Run Final Data Mapping**")
                if st.button("Run Data Mapping and Reconciliation", use_container_width=True):
                    try:
                        run_final_data_mapping_from_workspace(
                            result=result,
                            generated_target=generated_target,
                            legacy_data_dictionary_folder=legacy_data_dictionary_folder,
                            target_data_dictionary_folder=target_data_dictionary_folder,
                        )
                    except Exception as exc:
                        st.session_state["data_mapping_error"] = str(exc)
                        st.rerun()
            if st.session_state.get("data_mapping_error"):
                st.error(st.session_state["data_mapping_error"])
            data_mapping_result = st.session_state.get("data_mapping_result") or result.get("data_mapping_result")
            render_data_mapping_result(data_mapping_result)

    elif active_section == "Validation":
        if requirements_status != SME_APPROVED:
            st.info("Validation is blocked until the requirements draft is SME approved.")
        elif technical_spec_status != ARCHITECT_APPROVED:
            st.info("Validation is blocked until the technical specification is architect approved.")
        elif not result.get("forward_engineering_output"):
            st.info("Validation is available after forward engineering generates target artifacts.")
        else:
            if not (st.session_state.get("validation_result") or result.get("validation_result")):
                st.markdown("**Run Validation**")
                if st.button("Run Validation", use_container_width=True):
                    try:
                        run_validation_from_workspace(
                            result=result,
                            requirements_draft=requirements_draft,
                            technical_spec_draft=technical_spec_draft,
                            generated_target=generated_target,
                        )
                    except Exception as exc:
                        st.session_state["validation_error"] = str(exc)
                        st.rerun()
            if st.session_state.get("validation_error"):
                st.error(st.session_state["validation_error"])
            validation_result = st.session_state.get("validation_result") or result.get("validation_result")
            render_validation_result(validation_result)

    elif active_section == "AI Code Review":
        publish_result = st.session_state.get("github_publish_result") or result.get("github_publish_result") or {}
        publish_config = resolve_publish_config(result, target_repo_branch)
        if requirements_status != SME_APPROVED:
            st.info("AI Code Review is blocked until the requirements draft is SME approved.")
        elif technical_spec_status != ARCHITECT_APPROVED:
            st.info("AI Code Review is blocked until the technical specification is architect approved.")
        elif not result.get("forward_engineering_output"):
            st.info("AI Code Review is available after forward engineering generates target artifacts.")
        else:
            st.markdown("**GitHub PR Configuration**")
            st.caption(f"Repository: {target_repo_full_name}")
            if not get_github_token():
                st.warning("GITHUB_TOKEN is not configured. Set it in the environment before creating a PR or posting review comments.")

            if not st.session_state.get("github_base_branch"):
                st.session_state["github_base_branch"] = publish_config["base_branch"]
            if not st.session_state.get("github_branch_name"):
                st.session_state["github_branch_name"] = publish_config["branch_name"]
            if not st.session_state.get("github_commit_message"):
                st.session_state["github_commit_message"] = publish_config["commit_message"]
            if not st.session_state.get("github_pr_title"):
                st.session_state["github_pr_title"] = publish_config["pr_title"]

            base_col, branch_col = st.columns(2)
            with base_col:
                st.text_input("Base branch", key="github_base_branch")
            with branch_col:
                st.text_input("Publish branch", key="github_branch_name")

            st.text_input("Commit message", key="github_commit_message")
            st.text_input("PR title", key="github_pr_title")

            if st.session_state.get("github_publish_error"):
                st.error(st.session_state["github_publish_error"])
            if publish_result:
                render_table(
                    [
                        {"metric": "Repository", "value": publish_result.get("repository_full_name", "")},
                        {"metric": "Branch", "value": publish_result.get("branch_name", "")},
                        {"metric": "Base", "value": publish_result.get("base_branch", "")},
                        {"metric": "PR URL", "value": publish_result.get("pr_url", "") or "Created"},
                    ],
                    "No PR details available.",
                )

            if not (st.session_state.get("ai_code_review_result") or result.get("ai_code_review_result")):
                st.markdown("**Run AI Code Review**")
                st.caption("This review evaluates the generated code diff against approved requirements, technical design, validation, and mapping results. If no PR is stored in state, the app creates or recovers one from the current publish branch before posting the review.")
                if st.button("Create or Reuse PR and Post AI Review", use_container_width=True, disabled=not bool(get_github_token())):
                    try:
                        st.session_state["github_publish_error"] = None
                        st.session_state["ai_code_review_error"] = None
                        generated_diff = build_ai_code_review_input(forward_comparison_items)
                        with st.spinner("Generating AI code review findings and posting them to the PR..."):
                            if not publish_result:
                                current_publish_config = {
                                    "branch_name": st.session_state.get("github_branch_name", "").strip(),
                                    "base_branch": st.session_state.get("github_base_branch", "").strip(),
                                    "commit_message": st.session_state.get("github_commit_message", "").strip(),
                                    "pr_title": st.session_state.get("github_pr_title", "").strip(),
                                }
                                if not current_publish_config["branch_name"]:
                                    raise RuntimeError("No publish branch is available. Enter the branch and PR details in AI Code Review first.")
                                existing_pr = github_find_open_pr(
                                    target_repo_full_name,
                                    current_publish_config["branch_name"],
                                    current_publish_config["base_branch"],
                                )
                                if existing_pr:
                                    publish_result = {
                                        "repository_full_name": target_repo_full_name,
                                        "branch_name": current_publish_config["branch_name"],
                                        "base_branch": current_publish_config["base_branch"],
                                        "commit_message": current_publish_config["commit_message"],
                                        "commit_sha": "",
                                        "pr_title": existing_pr.get("title", current_publish_config["pr_title"]),
                                        "pr_number": existing_pr.get("number", 0),
                                        "pr_url": existing_pr.get("html_url", ""),
                                        "staged_files": [],
                                    }
                                else:
                                    publish_result = publish_generated_target_pr(
                                        generated_target=generated_target,
                                        result=result,
                                        comparison_items=forward_comparison_items,
                                        target_repo_full_name=target_repo_full_name,
                                        target_code_root=target_repo_code_root,
                                        branch_name=current_publish_config["branch_name"],
                                        base_branch=current_publish_config["base_branch"],
                                        commit_message=current_publish_config["commit_message"],
                                        pr_title=current_publish_config["pr_title"],
                                    )
                                st.session_state["github_publish_result"] = publish_result
                                st.session_state["analysis_result"]["github_publish_result"] = publish_result
                            review_result = ai_code_review_agent.run(
                                requirements=requirements_draft,
                                technical_spec=technical_spec_draft,
                                validation_result=st.session_state.get("validation_result") or result.get("validation_result") or {},
                                data_mapping_result=st.session_state.get("data_mapping_result") or result.get("data_mapping_result") or {},
                                generated_diff=generated_diff,
                                logger=lambda _a, _m: None,
                            )
                            github_post_pr_review_comment(
                                repo_full_name=publish_result.get("repository_full_name", ""),
                                pr_number=int(publish_result.get("pr_number", 0)),
                                review_body=review_result.get("review_comment_markdown", ""),
                            )
                        st.session_state["ai_code_review_result"] = review_result
                        st.session_state["analysis_result"]["ai_code_review_result"] = review_result
                        st.session_state["pending_main_section"] = "AI Code Review"
                        st.rerun()
                    except Exception as exc:
                        st.session_state["ai_code_review_error"] = str(exc)
                        st.rerun()
            if st.session_state.get("ai_code_review_error"):
                st.error(st.session_state["ai_code_review_error"])
            ai_code_review_result = st.session_state.get("ai_code_review_result") or result.get("ai_code_review_result")
            render_ai_code_review_result(ai_code_review_result)

    elif active_section == "Execution Logs":
        render_logs(result.get("logs", []))

    elif active_section == "Raw JSON":
        st.json(result, expanded=False)


if __name__ == "__main__":
    main()
